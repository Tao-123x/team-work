from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
DOWNLOADS = Path.home() / "Downloads"


TRANSLATIONS = {
    "TEAM PROJECT: Team Meeting Plan Template": "团队项目：团队会议纪要模板",
    "TEAM PROJECT: Team Leader Report Template": "团队项目：Team Leader 周报模板",
    "TEAM PROJECT: Project Problem Domain Activity Sheet": "团队项目：项目问题领域活动表",
    "TEAM PROJECT: Project Scoping & Planning Activity Sheet": "团队项目：项目范围与规划活动表",
    "The Team Leader of your group should open this file and update it every time you have a meeting and record the work allocated to each member.": "你们小组的 Team Leader 应打开此文件，并在每次开会后更新，记录分配给每位成员的工作。",
    "Ideally you should be holding stand-up meetings regularly. These review the previous activity and discuss what each team member will be doing on this day. The meeting may also discuss any challenges members are finding with tasks. The meeting should last around 15 minutes (and no more than 30). You should plan for at least 2 of the stand-ups to be in person (Physical), with the others held using MS Teams.": "理想情况下，你们应定期举行 stand-up 会议。会议用于回顾之前的工作，并讨论每位成员当天要做什么，也可以讨论成员在任务中遇到的困难。会议时长应控制在 15 分钟左右，最多不超过 30 分钟。你们应至少安排 2 次线下（Physical）stand-up，其余会议可通过 MS Teams 进行。",
    "You should also aim to have at least two allocated working sessions per week where all the team members share the same environment for a significant period of time (i.e. at least a couple of hours).": "你们还应尽量保证每周至少安排两次集中协作时段，让所有团队成员在同一环境中持续工作较长时间（例如至少几个小时）。",
    "Once updated save the file and upload it to your MS Teams chat group using the ‘Files’ area (preferable). If you are using Wechat, you can ‘Search Chat History’ and then filter by ‘Files’ to find out all your uploaded files in the group chat.": "更新完成后，请保存文件，并通过 MS Teams 群聊中的“Files”区域上传（推荐）。如果你们使用 WeChat，可以通过“搜索聊天记录”并按“文件”筛选，找到群聊中已上传的所有文件。",
    "Once completed do not forget to save your work as [00_ Team Meeting Plan Template].": "完成后，请不要忘记将文件保存为 [00_ Team Meeting Plan Template]。",
    "Upload the completed report to your Microsoft Teams ‘Files’ area for the MS Teams channel you have been provided with.": "请将完成后的报告上传到分配给你们的 MS Teams 频道中的 “Files” 区域。",
    "As the Team Leader, you should update this file on a regular basis and ensure it is up to date following each meeting regardless of virtual or in person.": "作为 Team Leader，你应定期更新此文件，并确保每次会议后都及时同步内容，无论会议是线上还是线下进行。",
    "LINE MANAGER: \n[insert tutors name]": "LINE MANAGER:\n[insert tutors name]",
    "GROUP NAME:": "GROUP NAME：",
    "TEAM COMPOSITION": "团队成员构成",
    "ROLE": "角色",
    "FULL NAME": "姓名",
    "STUDENT ID": "学号",
    "Team Leader": "Team Leader（团队负责人）",
    "Product Owner": "Product Owner（产品负责人）",
    "Quality Controller": "Quality Controller（质量控制员）",
    "Team Member": "Team Member（团队成员）",
    "MEETING SUMMARY": "会议纪要",
    "Add more rows as needed": "如有需要可增加更多行",
    "Date / Time": "日期 / 时间",
    "Virtual / Physical (enter location)": "线上 / 线下（填写地点）",
    "Duration": "时长",
    "Meeting Summary": "会议摘要",
    "Member(s) Present": "出席成员",
    "Example": "示例",
    "Physical": "线下",
    "30 mins": "30分钟",
    "Group roles assigned \nBegan brainstorming project ideas": "已分配小组角色\n开始头脑风暴项目想法",
    "All members": "全体成员",
    "The Team Leader of your group should open this file and update it each week.": "你们小组的 Team Leader 应打开此文件，并在每周进行更新。",
    "The Team Leader must ensure this file is updated regularly to review with the Line Manager every week.": "Team Leader 必须确保此文件定期更新，以便每周与 Line Manager 一起回顾。",
    "Once updated save the file and upload to your MS Teams chat group using the ‘Files’ area.": "更新完成后，请保存文件，并通过 MS Teams 群聊中的“Files”区域上传。",
    "Once completed do not forget to save your work as [00_ Team Leader Report Template]": "完成后，请不要忘记将文件保存为 [00_ Team Leader Report Template]。",
    "As the Team Leader you should update this file on a regular basis and ensure it is up to date ahead of your Bi-Weekly meetings with your Line Manager.": "作为 Team Leader，你应定期更新此文件，并确保在与你的 Line Manager 进行双周会议前内容已经更新完成。",
    "GROUP NAME:\n[allocated by Line Manager]": "GROUP NAME:\n[allocated by Line Manager]",
    "PART 1 - PRODUCT PROGRESS": "PART 1 - 产品进展",
    "In this page, you should give each-session summary of the project’s progress. This should highlight what has been done, and needs to be done, with an indication as to whether the activity is on schedule to deliver the project.": "在本页中，你应按每次课/每次会议总结项目进展。内容应突出已经完成的事项、尚待完成的事项，以及当前活动是否按计划推进并能够支持项目按时交付。",
    "Week Number": "周次",
    "Tasks Done": "已完成任务",
    "Tasks Pending": "待完成任务",
    "Project Schedule (Ahead/ Ontime/ Behind)": "项目进度（提前 / 按时 / 落后）",
    "PART 2 - TEAM ENGAGEMENT": "PART 2 - 团队参与情况",
    "In this section, you should report on the weekly team engagement. As this is an open document you should avoid naming specific people, but you should review how the team are engaging with the activity and how you are working with the team to ensure they are motivated and clear about what they are doing. \nIf there are specific individuals of concern these should be raised with your Line Manager.": "在本部分中，你应汇报团队每周的参与情况。由于这是开放文档，请避免点名具体成员；你应评述团队整体如何参与各项活动，以及你如何与团队协作以确保成员保持积极性并清楚自己正在做什么。\n如果有需要特别关注的个体，应单独向 Line Manager 说明。",
    "Team Engagement Review": "团队参与情况评述",
    "Week 2": "Week 2（第2周）",
    "Week 3": "Week 3（第3周）",
    "Week 4": "Week 4（第4周）",
    "Week 5": "Week 5（第5周）",
    "Week 6": "Week 6（第6周）",
    "PART 3 -  MANAGEMENT ISSUES": "PART 3 - 管理问题",
    "In this section, you should outline any management issues you have or problems you have overcome each week.  As this is an open document you should avoid naming specific people. If there are specific individuals of concern these should be raised with your Line Manager.": "在本部分中，你应概述每周遇到的管理问题，或本周已经克服的问题。由于这是开放文档，请避免点名具体成员。如果有需要特别关注的个体，应单独向 Line Manager 说明。",
    "The Product Owner of your group should open this file and all group members should contribute to completing this activity sheet.": "你们小组的 Product Owner 应打开此文件，并由全组成员共同参与完成这份活动表。",
    "Once completed save the file and upload it to your MS Teams chat group using the ‘Files’ area.": "完成后，请保存文件，并通过 MS Teams 群聊中的“Files”区域上传。",
    "Once completed do not forget to save your work as [02_ Project Scoping & Planning Activity Sheet]": "完成后，请不要忘记将文件保存为 [02_ Project Scoping & Planning Activity Sheet]。",
    "As a team you should refer back to this document during this sprint to ensure your emerging product remains focused on the planning and goals.": "作为团队，你们应在本次 sprint 期间持续回看此文档，确保正在形成的产品始终围绕既定计划和目标推进。",
    "This document will become useful when you come to make your team presentations. The better you document the planning of your project now, the easier it will be to reflect on your progress at the end.": "当你们后续准备团队展示时，这份文档会很有帮助。现在把项目规划记录得越清楚，最后回顾项目进展时就会越容易。",
    "PROBLEM DOMAIN": "问题领域",
    "Insert the Problem Domain that you have agreed to work on from Week 1 Activity Sheet to assist you with the completion of this activity": "请填入你们在 Week 1 Activity Sheet 中已经确定的问题领域，以辅助完成本次活动。",
    "What is the actual problem?": "实际问题是什么？",
    "Why are you doing it?": "你们为什么要做这个？",
    "How are you going to go about addressing it?": "你们打算如何推进并解决这个问题？",
    "INITIAL CONCEPT": "初始概念",
    "Give a brief description of your initial concept. If you have more than one you should use a separate table for each (but try to stick to one).": "请简要描述你们的初始概念。如果你们有多个概念，应为每个概念使用单独的表格（但尽量控制为一个）。",
    "Briefly document the initial concept in high-level terms that describe the type of solution and how it addresses the ‘What’ and ‘Why’ of the problem domain. (This could be in the form of a ‘user story’ or ‘use case’)": "请用高层次的方式简要记录初始概念，说明它属于什么类型的解决方案，以及它如何回应问题领域中的“做什么”和“为什么做”。（可以采用“user story”或“use case”的形式）",
    "Briefly document the 3 key features of this concept. \n(Avoid thinking about technology, focus on what they do. These will become the focus of your project ‘sprints’)": "请简要记录该概念的 3 个关键功能。\n（避免先考虑技术实现，重点写它们“做什么”。这些功能将成为你们项目各次 sprint 的重点。）",
    "Briefly document the 3 key features of this concept.": "请简要记录该概念的 3 个关键功能。",
    "(Avoid thinking about technology, focus on what they do. These will become the focus of your project ‘sprints’)": "（避免先考虑技术实现，重点写它们“做什么”。这些功能将成为你们项目各次 sprint 的重点。）",
    "Feature 1:": "功能 1：",
    "Feature 2:": "功能 2：",
    "Feature 3:": "功能 3：",
    "SOLUTION CONCEPT": "解决方案概念",
    "Expand on your initial concept (Features 1, 2, 3) and start to define what this solution might be. You should consider 3 developmental stages (Sprints 1, 2,3) and describe, in a single sentence, what they would be. For the stages it is a good idea to think about what the system/solution might ‘look’ like at the end of each stage.": "请在初始概念（功能 1、2、3）的基础上进一步展开，开始定义该解决方案可能会是什么样子。你们应考虑 3 个开发阶段（Sprint 1、2、3），并分别用一句话描述每个阶段的成果。思考这些阶段时，最好设想系统/方案在每个阶段结束时“看起来会是什么样”。",
    "Development Stage 1 for product": "产品开发阶段 1",
    "Development Stage 2 for product": "产品开发阶段 2",
    "Development Stage 3 for product": "产品开发阶段 3",
    "OUTLINE PLAN": "概要计划",
    "This section will focus on a breakdown of how you plan on delivering the project using 3 Sprints (Sprints 1, 2, 3). Use the table below to plan your First Sprint. Use your Trello Board to plan this.": "本部分聚焦于你们如何通过 3 个 Sprint（Sprint 1、2、3）来推进项目交付。请使用下表规划你们的第一个 Sprint，并结合你们的 Trello 看板进行计划。",
    "SPRINT 1: PLANNING": "SPRINT 1：规划",
    "This is the activity that will take place between Week 2, Week 3, and Week 4 sessions. It should be mostly about research, planning, and the first stage of the development of your product (the activity for this sprint has been partially filled in for you to get you started).": "这是发生在 Week 2、Week 3 和 Week 4 课程之间的活动。内容应主要围绕研究、规划，以及产品开发的第一阶段展开（本次 sprint 的部分内容已为你们预填，帮助你们开始）。",
    "Activity:": "活动：",
    "Research and prototype aspects of the solution, define project tasks, and populate Trello board. Start to develop structure/frameworks/architecture components for the solution.": "研究并原型化解决方案的相关方面，定义项目任务，并完善 Trello 看板。开始搭建解决方案的结构 / 框架 / 架构组件。",
    "Objective: \nYou should define an objective for the end of the week (i.e. what would a successful week look like).": "目标：\n你们应定义本周结束时要达成的目标（即怎样才算度过了成功的一周）。",
    "Objective:": "目标：",
    "You should define an objective for the end of the week (i.e. what would a successful week look like).": "你们应定义本周结束时要达成的目标（即怎样才算度过了成功的一周）。",
    "Goals\nYou should consider a set of goals for the sprint and estimate how many person-hours they will take to deliver (you will get better at time estimation as this process goes, so don’t worry too much about being wrong here). Think about 5-6 goals for this sprint. Try not to overly define them, but be realistic about time.": "Goals（目标）\n你们应为本次 sprint 设定一组目标，并估算完成它们需要多少人小时（随着过程推进，你们会越来越擅长时间估算，所以现在估错一点也没关系）。本次 sprint 可以考虑 5-6 个目标。尽量不要定义得过细，但时间估算要尽量现实。",
    "Goals": "Goals（目标）",
    "You should consider a set of goals for the sprint and estimate how many person-hours they will take to deliver (you will get better at time estimation as this process goes, so don’t worry too much about being wrong here). Think about 5-6 goals for this sprint. Try not to overly define them, but be realistic about time.": "你们应为本次 sprint 设定一组目标，并估算完成它们需要多少人小时（随着过程推进，你们会越来越擅长时间估算，所以现在估错一点也没关系）。本次 sprint 可以考虑 5-6 个目标。尽量不要定义得过细，但时间估算要尽量现实。",
    "Goal 1:": "目标 1：",
    "Goal 2:": "目标 2：",
    "Goal 3:": "目标 3：",
    "Goal 4:": "目标 4：",
    "Goal 5:": "目标 5：",
    "Objective": "目标",
    "Time Estimate (hours)": "时间估算（小时）",
    "The Total time estimate is the sum of the time estimates for each goal. It should reflect 10 hours per person to be completed between this session and next week, as part of your group activity.": "总时间估算是各个目标时间估算的总和。它应反映从本次课到下周之间，每位成员大约投入 10 小时的团队活动工作量。",
    "Total Time Estimate (hours)": "总时间估算（小时）",
    "The Team Leader of your group should open this file and all group members should contribute to completing this activity sheet.": "你们小组的 Team Leader 应打开此文件，并由全组成员共同参与完成这份活动表。",
    "Once completed save the file and upload it to your MS Teams chat group using the ‘Files’ area (preferable). If you are using WeChat, you can ‘Search Chat History’ and then filter by ‘Files’ to find out all your uploaded files in the group. chat.": "完成后，请保存文件，并通过 MS Teams 群聊中的“Files”区域上传（推荐）。如果你们使用 WeChat，可以通过“搜索聊天记录”并按“文件”筛选，找到群聊中已上传的所有文件。",
    "Once completed do not forget to save your work as [01_ Project Problem Domain Activity Sheet]": "完成后，请不要忘记将文件保存为 [01_ Project Problem Domain Activity Sheet]。",
    "Upload the completed report to your Microsoft Teams ‘Files’ area for the MS Teams chat group you have been provided with (preferable). If you are using WeChat, you can ‘Search Chat History’ and then filter by ‘Files’ to find out all your uploaded files in the group chat.": "请将完成后的报告上传到分配给你们的 MS Teams 群聊中的 “Files” 区域（推荐）。如果你们使用 WeChat，可以通过“搜索聊天记录”并按“文件”筛选，找到群聊中已上传的所有文件。",
    "As a team you should discuss these initial ideas and agree on a final product to work on for the project.": "作为团队，你们应讨论这些初始想法，并确定最终要开展的项目产品。",
    "Use this space to brainstorm some possible project ideas. Use a new table for each new idea. This should define the high-level problem, without thinking of any type of solution and identify the 3 key aspects of the problem domain.": "请使用此区域头脑风暴一些可能的项目想法。每产生一个新想法，就使用一张新的表格。这里应定义高层次的问题本身，而不是先考虑解决方案，并识别该问题领域的 3 个关键方面。",
    "What is the actual problem?\n(What you do aim to do)": "实际问题是什么？\n（你们想要做成什么）",
    "Why are you doing it?\n(What is the benefit of doing this and/or the impact of not doing it)": "你们为什么要做这个？\n（这样做的好处是什么，或者不做会带来什么影响）",
    "How are you going to go about addressing it? \n(This is the most difficult part. You should not be thinking about the solution, but the process by which you are going to research, design/implement, and evaluate and the type of potential solution to this problem domain. It should be high level and strategic.)": "你们打算如何着手解决这个问题？\n（这是最难的一部分。你们不应直接考虑具体解决方案，而应考虑你们将如何进行研究、设计 / 实现、评估，以及适用于该问题领域的潜在解决方案类型。这里应保持高层次和战略性。）",
}


SOURCES = [
    (
        DOWNLOADS / "Team Meeting Summary Template.docx",
        ROOT / "团队会议纪要模板_中文版.docx",
    ),
    (
        DOWNLOADS / "Team Leader Report Template.docx",
        ROOT / "团队负责人周报模板_中文版.docx",
    ),
    (
        DOWNLOADS / "Project Problem Domain Activity Sheet.docx",
        ROOT / "项目问题领域活动表_中文版.docx",
    ),
    (
        DOWNLOADS / "Project Scoping  Planning Activity Sheet.docx",
        ROOT / "项目范围与规划活动表_中文版.docx",
    ),
]


def translate_text(text: str) -> str:
    stripped = text.strip()
    if stripped in TRANSLATIONS:
        return TRANSLATIONS[stripped]
    return TRANSLATIONS.get(text, text)


def translate_paragraphs(paragraphs) -> None:
    for paragraph in paragraphs:
        original = paragraph.text
        translated = translate_text(original)
        if translated != original:
            paragraph.text = translated
        apply_font(paragraph)


def translate_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                original = paragraph.text
                translated = translate_text(original)
                if translated != original:
                    paragraph.text = translated
                apply_font(paragraph)
            for nested_table in cell.tables:
                translate_table(nested_table)


def apply_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Songti SC"
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "Songti SC")
        r_fonts.set(qn("w:ascii"), "Songti SC")
        r_fonts.set(qn("w:hAnsi"), "Songti SC")


def translate_document(source: Path, target: Path) -> None:
    doc = Document(str(source))
    translate_paragraphs(doc.paragraphs)
    for table in doc.tables:
        translate_table(table)
    for section in doc.sections:
        translate_paragraphs(section.header.paragraphs)
        translate_paragraphs(section.footer.paragraphs)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))


def main() -> None:
    for source, target in SOURCES:
        translate_document(source, target)
        print(f"Created: {target}")


if __name__ == "__main__":
    main()
