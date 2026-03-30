from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials" / "week4"

TEAM = [
    ("Jiacheng Tao", "Team Leader"),
    ("Pengyang Liu", "Product Owner"),
    ("Sicheng Wu", "Quality Controller"),
    ("Xiaofeng Dang", "Team Member (Development)"),
    ("Ruilin Ma", "Team Member (Development)"),
]

WORKSTREAMS = [
    (
        "Lifecycle completion",
        "Confirm that the main order journey now includes accept, picked up, delivered, completed, and cancelled states.",
    ),
    (
        "Runtime verification",
        "Check whether the current build can prove the lifecycle works through controlled runtime evidence.",
    ),
    (
        "Data reset and repeatability",
        "Review whether database reset and seeded data make the demo path repeatable for later rehearsals.",
    ),
    (
        "Week 5 readiness",
        "Decide what still needs to improve for demo stability and presentation confidence.",
    ),
]

POST_CRITERIA_NOTES = [
    (
        "Lifecycle Coverage",
        "Does the build now support the complete core order path?",
        "The team checked whether the relay workflow had moved beyond partial progress into a full, defensible order lifecycle.",
    ),
    (
        "Verification Honesty",
        "Can the team prove the core path with believable runtime evidence?",
        "The group wanted not just code progress, but evidence that the workflow could be rerun and observed honestly.",
    ),
    (
        "Repeatability",
        "Can the same scenario be reset and demonstrated again?",
        "Reset support mattered because the final demo would need predictable seeded data and controlled state.",
    ),
    (
        "Presentation Impact",
        "Does this progress reduce risk for the final defense?",
        "The team evaluated Week 4 not only as implementation progress, but as a step toward a stable classroom presentation.",
    ),
]

POST_DECISIONS = [
    ("Main Technical Outcome", "The core order lifecycle was treated as the central Week 4 milestone and brought much closer to a complete runnable flow."),
    ("Verification Outcome", "The team confirmed that runtime verification and repeatable reset paths were now necessary evidence, not optional extras."),
    ("Next Sprint Direction", "Week 5 would focus on demo stability, smoother navigation, and presentation-oriented polish around the verified core path."),
]

POST_FOLLOW_UP = [
    "Product Owner to keep Week 5 focused on a controllable demo path rather than adding broad new features.",
    "Team Leader to turn the verified lifecycle into a clear presentation checkpoint and next-step action list.",
    "Quality Controller to formalize the runtime proof and identify any remaining weak spots in the verified path.",
    "Developers to improve weak transitions, support reset-ready demos, and strengthen the same tested user journey.",
]

POST_DISCUSSION_SUMMARY = [
    (
        "What changed in Week 4",
        "The project moved from early integration into a more complete and testable order lifecycle, which made the product story much stronger.",
    ),
    (
        "Why verification mattered",
        "The team agreed that a feature only really counted once it could be demonstrated through repeatable runtime evidence.",
    ),
    (
        "How this affected the project",
        "By locking onto the full lifecycle, the group created a stable base for later demo polishing instead of continuing to build on uncertain behavior.",
    ),
]

POST_MEMBER_NOTES = {
    "Jiacheng Tao": [
        "I focused on whether Week 4 could convert previous integration work into one verified and manageable story for the team.",
        "I pushed the group to discuss evidence and repeatability, not only feature completion, because that would matter in the defense.",
        "I agreed that finishing the lifecycle was the right milestone only if it could also reduce future demo risk.",
        "I left the meeting with responsibility for aligning the verified build progress with the next presentation-oriented sprint.",
    ],
    "Pengyang Liu": [
        "I focused on whether the completed lifecycle still matched the original MVP story without becoming bloated.",
        "I argued that the core path should remain easy to explain even as more statuses and transitions were added.",
        "I agreed that runtime proof was important because it showed the product flow was not just planned but genuinely usable.",
        "I took away the task of preserving the clear relay narrative while preparing the next sprint's priorities.",
    ],
    "Sicheng Wu": [
        "I focused on whether the team could now prove the lifecycle honestly and repeatedly instead of relying on assumptions.",
        "I highlighted that reset support and sequential verification were essential to make future demo evidence credible.",
        "I agreed that a stronger verification story would raise both build-quality marks and presentation confidence.",
        "I left the meeting planning to turn observed runtime checks into a clearer QA-facing record for the next stage.",
    ],
    "Xiaofeng Dang": [
        "I focused on whether the now-larger order flow still felt understandable from the user's point of view.",
        "I suggested that even if the backend path worked, the interface still had to make each stage readable during demonstration.",
        "I agreed that verified functionality needed equally clear front-end guidance if the audience was going to follow the story.",
        "I took away the need to support the verified path with cleaner interactions and more stable presentation flow.",
    ],
    "Ruilin Ma": [
        "I focused on whether the order states now formed a complete and reliable logic chain rather than a set of isolated actions.",
        "I pointed out that repeatable reset and lifecycle verification were signs that the backend was maturing into something dependable.",
        "I agreed that the team should protect this verified path instead of destabilizing it with unnecessary changes.",
        "I left the meeting clearer on which logic and stability improvements would matter most before the demo-oriented sprint.",
    ],
}

POST_TL_NOTES = [
    "Week 4 was the point where implementation progress became much more defensible because the team centered discussion on the full order lifecycle.",
    "The meeting helped the group distinguish between code that existed and code that had actually been verified through a repeatable scenario.",
    "By the end of the review, the project had a much stronger backbone for both technical explanation and controlled demonstration.",
    "The meeting also prevented scope drift by directing Week 5 toward demo stability instead of new ambitious additions.",
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def write_cell(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(10.5)
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_meta(doc: Document, status_label: str, pre: bool) -> None:
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if pre:
        rows = [
            ("Week", "Week 4", "Document Status", status_label),
            ("Meeting Type", "Lifecycle Completion and Verification Review", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "__________________", "Location", "__________________"),
            ("Attendees", "All group members", "Recorder", "__________________"),
        ]
    else:
        rows = [
            ("Week", "Week 4", "Document Status", status_label),
            ("Meeting Type", "Lifecycle Completion and Verification Review", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "Use the actual Week 4 meeting time from the handwritten archive copy.", "Location", "Use the actual meeting location from the handwritten archive copy."),
            ("Attendees", "All group members", "Recorder", "Use the recorder written on the paper archive copy."),
        ]
    for row, values in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)


def add_participants(doc: Document, pre: bool) -> None:
    add_heading(doc, "Attendees and Roles")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Member", "Role"]
    if pre:
        headers.append("Personal Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for name, role in TEAM:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.52 if pre else 0.36)
        values = [name, role]
        if pre:
            values.append("\n" * 4)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_workstreams(doc: Document, pre: bool) -> None:
    add_heading(doc, "Workstreams Under Review")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Workstream", "Review Focus"]
    if pre:
        headers.append("Meeting Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for workstream, note in WORKSTREAMS:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.68 if pre else 0.45)
        values = [workstream, note]
        if pre:
            values.append("\n" * 5)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_criteria(doc: Document, pre: bool) -> None:
    add_heading(doc, "Review Criteria")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Criteria", "Purpose", "Meeting Record"]):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    criteria = [
        ("Lifecycle Coverage", "Does the build now support the complete core order path?"),
        ("Verification Honesty", "Can the team prove the core path with believable runtime evidence?"),
        ("Repeatability", "Can the same scenario be reset and demonstrated again?"),
        ("Presentation Impact", "Does this progress reduce risk for the final defense?"),
    ]
    for idx, (label, purpose) in enumerate(criteria):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.78 if pre else 0.56)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], purpose, size=10)
        if pre:
            write_cell(row.cells[2], "\n" * 5, size=10)
        else:
            write_cell(row.cells[2], POST_CRITERIA_NOTES[idx][2], size=10)


def add_large_notes_table(doc: Document, title: str, rows: list[str], height: float) -> None:
    add_heading(doc, title)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, label in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(height)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], "\n" * 7, size=10)


def add_post_summary(doc: Document) -> None:
    add_heading(doc, "Meeting Decisions")
    table = doc.add_table(rows=len(POST_DECISIONS), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DECISIONS):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.48)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)

    add_heading(doc, "Week 5 Follow-Up Actions")
    for item in POST_FOLLOW_UP:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Post-Meeting Discussion Summary")
    table = doc.add_table(rows=len(POST_DISCUSSION_SUMMARY), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DISCUSSION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.82)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)


def add_pre_body(doc: Document) -> None:
    add_title(doc, "Week 4 Lifecycle Completion and Verification Review", "Pre-Meeting Discussion Base")
    add_meta(doc, "Pre-meeting", pre=True)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review whether the relay workflow has now reached a complete order lifecycle.",
        "Check what runtime evidence already exists and what still needs to be verified honestly.",
        "Decide how Week 5 should build on a repeatable, presentation-safe core path.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=True)
    add_workstreams(doc, pre=True)
    add_criteria(doc, pre=True)
    add_large_notes_table(
        doc,
        "Open Discussion Space",
        [
            "What parts of the lifecycle now look complete",
            "What still feels weak, risky, or not yet well verified",
            "What evidence must be captured or improved before the next sprint ends",
        ],
        1.05,
    )
    add_large_notes_table(
        doc,
        "Decision and Next Actions (leave blank before the meeting)",
        [
            "Main lifecycle conclusions agreed in the meeting",
            "Verification actions selected for the next sprint",
            "Actions to complete before Week 5 review ends",
        ],
        1.0,
    )


def add_post_body(doc: Document) -> None:
    add_title(doc, "Week 4 Lifecycle Completion and Verification Review", "Post-Meeting Archive Reference")
    add_meta(doc, "Post-meeting", pre=False)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review whether the relay workflow has now reached a complete order lifecycle.",
        "Check what runtime evidence already exists and what still needs to be verified honestly.",
        "Decide how Week 5 should build on a repeatable, presentation-safe core path.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=False)
    add_workstreams(doc, pre=False)
    add_criteria(doc, pre=False)
    add_post_summary(doc)


def add_individual_page(doc: Document, member: tuple[str, str], pre: bool) -> None:
    name, role = member
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        f"Week 4 Individual Notes - {name}",
        "Pre-Meeting Personal Discussion Sheet" if pre else "Post-Meeting Personal Record",
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name", name, "Role", role),
        ("Meeting Focus", "Lifecycle completion and runtime verification", "Week", "Week 4"),
    ]
    for row, values in zip(table.rows, rows):
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)

    if pre:
        add_large_notes_table(
            doc,
            "Personal Handwritten Space",
            [
                "What I most want to verify in the current lifecycle",
                "What still worries me about stability or evidence",
                "What progress I think is most important this week",
                "What I need clarified before the meeting ends",
            ],
            1.05,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Post-Meeting Personal Record")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "What I focused on during the lifecycle review",
            "What I said or suggested",
            "What I agreed or disagreed with",
            "What task or next step I took away",
        ]
        for row, label, note in zip(table.rows, labels, POST_MEMBER_NOTES[name]):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def add_tl_page(doc: Document, pre: bool) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        "Week 4 Team Leader Pre-Meeting Control Page" if pre else "Week 4 Team Leader Official Record",
        "Meeting setup and observation sheet" if pre else "Formal team archive version",
    )
    if pre:
        add_large_notes_table(
            doc,
            "Team Leader Writing Space",
            [
                "What must be verified honestly in this meeting",
                "How to separate real proof from optimistic assumptions",
                "What decisions must be fixed before the meeting ends",
                "How to keep Week 5 focused on stability rather than new scope",
            ],
            1.08,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Team Leader Official Summary")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "How Week 4 progress was evaluated",
            "Why verification became a core meeting topic",
            "How the team protected the next sprint from scope drift",
            "Leader reflection on the meeting quality",
        ]
        for row, label, note in zip(table.rows, labels, POST_TL_NOTES):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def create_doc(filename: str, pre: bool, member: tuple[str, str] | None = None, tl: bool = False) -> Path:
    doc = Document()
    set_defaults(doc)
    if pre:
        add_pre_body(doc)
    else:
        add_post_body(doc)

    if member:
        add_individual_page(doc, member, pre=pre)
    if tl:
        add_tl_page(doc, pre=pre)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_doc("Week4_PreMeeting_Shared_Base.docx", pre=True),
        create_doc("Week4_PostMeeting_Shared_Record.docx", pre=False),
        create_doc("Week4_PreMeeting_Team_Leader_Page.docx", pre=True, tl=True),
        create_doc("Week4_PostMeeting_Team_Leader_Official_Record.docx", pre=False, tl=True),
    ]

    for member in TEAM:
        safe = member[0].replace(" ", "_")
        created.append(create_doc(f"Week4_PreMeeting_{safe}_Notes.docx", pre=True, member=member))
        created.append(create_doc(f"Week4_PostMeeting_{safe}_Notes.docx", pre=False, member=member))

    print("Created files:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
