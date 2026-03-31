from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials"

SECTIONS = [
    (
        "Team Leader: Jiacheng Tao",
        [
            ("Meeting 1", [
                "Need fair comparison across all 5 ideas",
                "FlowRead interesting, but team scope may become wide",
                "weather push easier to explain in short presentation",
                "must decide role structure early",
                "next: confirm tools + planning session",
            ]),
            ("Meeting 2", [
                "tutor feedback means old topic not strong enough",
                "pivot must be controlled, not random",
                "takeaway relay has clearer user pain point",
                "need backlog reset immediately",
                "next: assign owners + keep everyone aligned",
            ]),
            ("Meeting 3", [
                "check if everyone is building toward same MVP",
                "integration gaps > adding features now",
                "need one issue list, not scattered comments",
                "Week 4 should focus on core flow only",
                "next: track blockers + priorities",
            ]),
            ("Meeting 4", [
                "full lifecycle is key milestone",
                "evidence matters, not just code exists",
                "reset path needed for repeatable demo",
                "protect verified flow from risky extra changes",
                "next: connect runtime proof to presentation story",
            ]),
            ("Meeting 5", [
                "final meeting must lock everything",
                "QR access + speaking order + paper archive all linked",
                "no more scope expansion before defense",
                "prepare backup plan if live demo unstable",
                "next: finalize run sheet + archive order",
            ]),
        ],
    ),
    (
        "Product Owner: Pengyang Liu",
        [
            ("Meeting 1", [
                "compare projects by user value, not only interest",
                "broad platform ideas may weaken final MVP",
                "weather push easier to frame as one clear problem",
                "need clearer scope boundary from start",
                "next: draft first backlog",
            ]),
            ("Meeting 2", [
                "new topic must answer tutor feedback directly",
                "dorm takeaway problem is more concrete",
                "relay workflow can stay simple: post -> accept -> deliver",
                "avoid feature expansion during pivot",
                "next: rewrite backlog around new user story",
            ]),
            ("Meeting 3", [
                "is current build still matching MVP story?",
                "main journey still needs smoothing",
                "weak details can wait, core path first",
                "need clearer user-facing progression",
                "next: sharpen Week 4 priorities",
            ]),
            ("Meeting 4", [
                "full lifecycle now supports product story better",
                "status chain easier to explain now",
                "must keep explanation simple",
                "verified path > extra features",
                "next: preserve one clear relay narrative",
            ]),
            ("Meeting 5", [
                "defense should show one short happy path only",
                "audience should understand product from first screen",
                "don't overload demo with exceptions",
                "printed evidence should support teamwork story",
                "next: keep explanation narrow + user-centered",
            ]),
        ],
    ),
    (
        "Quality Controller: Sicheng Wu",
        [
            ("Meeting 1", [
                "which proposal is easiest to evaluate honestly?",
                "need records from the beginning",
                "some ideas too broad -> hard to test well",
                "must think about evidence early",
                "next: prepare first checklist",
            ]),
            ("Meeting 2", [
                "pivot means quality plan also changes",
                "new topic easier to test with role-based flow",
                "need fresh acceptance criteria",
                "must record tutor-feedback impact clearly",
                "next: rebuild checklist for new MVP",
            ]),
            ("Meeting 3", [
                "what bugs already visible?",
                "integration issues should be discussed now",
                "check confusing states / weak transitions",
                "evidence of real review is important",
                "next: convert issues into QA checklist",
            ]),
            ("Meeting 4", [
                "can we prove lifecycle honestly and repeatedly?",
                "reset support is part of quality",
                "sequential verification safer than assumptions",
                "completed / cancelled proof strengthens credibility",
                "next: keep runtime evidence organized",
            ]),
            ("Meeting 5", [
                "final path must stay verified",
                "backup plan needed if live step fails",
                "no untested late changes now",
                "printed material also part of quality evidence",
                "next: prepare final readiness checklist",
            ]),
        ],
    ),
    (
        "Development Member: Xiaofeng Dang",
        [
            ("Meeting 1", [
                "which idea has clearest UI story?",
                "audience must understand value quickly",
                "weather push easier to demo visually",
                "page flow should not be too complex",
                "next: think about early screen flow",
            ]),
            ("Meeting 2", [
                "relay idea gives clearer requester/helper interaction",
                "better for visible user journey",
                "can show task ownership more clearly",
                "must keep pages simple after pivot",
                "next: think about posting / list / detail screens",
            ]),
            ("Meeting 3", [
                "some interactions still not smooth enough",
                "need clearer route between pages",
                "UI should support main workflow",
                "detail / list relationship needs attention",
                "next: improve visible flow points",
            ]),
            ("Meeting 4", [
                "verified backend flow needs readable interface too",
                "each lifecycle stage should be obvious in UI",
                "if audience gets lost, good logic won't help",
                "support main tested flow first",
                "next: strengthen demo screens",
            ]),
            ("Meeting 5", [
                "only keep cleanest visible path for defense",
                "no awkward jumps between screens",
                "speaking and UI should match",
                "protect visible pages from last-minute risk",
                "next: polish story-critical interactions",
            ]),
        ],
    ),
    (
        "Development Member: Ruilin Ma",
        [
            ("Meeting 1", [
                "which idea has most controllable logic?",
                "some concepts attractive, but backend may get messy",
                "easier state path means safer MVP",
                "weather push simpler at this stage",
                "next: think about technical structure",
            ]),
            ("Meeting 2", [
                "takeaway relay gives clearer state transitions",
                "requester / helper roles easier to model",
                "better fit for controlled backend path",
                "can structure data around order lifecycle",
                "next: focus on routes + order states",
            ]),
            ("Meeting 3", [
                "are current logic pieces connecting well?",
                "weak integration will hurt later if ignored now",
                "must protect core data / state flow",
                "not time for extra complexity",
                "next: fix main logic gaps first",
            ]),
            ("Meeting 4", [
                "lifecycle now looks more complete",
                "reset + repeatability show backend is maturing",
                "need to protect stable state path",
                "verified transitions > extra additions",
                "next: improve stability around tested flow",
            ]),
            ("Meeting 5", [
                "locked demo path must still match real runtime state",
                "shortcuts only safe if state stays consistent",
                "avoid risky backend changes before defense",
                "deployment and demo prep now linked",
                "next: keep runtime path stable",
            ]),
        ],
    ),
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_run(paragraph, text: str, *, bold: bool = False, size: int = 11, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "Handwriting Examples", bold=True, size=20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "Five members across the five agreed key meetings", italic=True, size=11)

    intro = doc.add_paragraph()
    add_run(intro, "Use these as realistic examples, not as text to copy identically. Keep each handwritten page to 5 to 8 short lines.", size=11)

    for section_title, meetings in SECTIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        add_run(p, section_title, bold=True, size=14)

        for meeting_title, bullets in meetings:
            mp = doc.add_paragraph()
            mp.paragraph_format.space_before = Pt(4)
            add_run(mp, meeting_title, bold=True, size=12)
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                add_run(bp, bullet, size=10.5)

    tips = doc.add_paragraph()
    tips.paragraph_format.space_before = Pt(10)
    add_run(tips, "Variation Tips", bold=True, size=14)

    for bullet in [
        "Mix short English phrases with a few Chinese notes if natural.",
        "Use arrows, underlines, and circles to make the writing look like real meeting notes.",
        "Do not let all five members write the same structure or wording.",
        "Prioritize Meeting 1, Meeting 2, and Meeting 5 if time is tight.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, bullet, size=10.5)

    path = OUTPUT_DIR / "Handwriting_Examples.docx"
    doc.save(path)
    print(path)


if __name__ == "__main__":
    main()
