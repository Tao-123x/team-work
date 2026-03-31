from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials"

TEAM = [
    ("Jiacheng Tao", "Team Leader"),
    ("Pengyang Liu", "Product Owner"),
    ("Sicheng Wu", "Quality Controller"),
    ("Xiaofeng Dang", "Development Member"),
    ("Ruilin Ma", "Development Member"),
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def write_run(paragraph, text: str, *, bold: bool = False, size: int = 11, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def write_cell(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    write_run(p, text, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_center_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(p, title, bold=True, size=20)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(sp, subtitle, italic=True, size=11)


def build_cover() -> Path:
    doc = Document()
    set_defaults(doc)

    doc.add_paragraph()
    doc.add_paragraph()
    add_center_title(
        doc,
        "Group 15 Teamwork Evidence Archive",
        "Dorm Takeaway Relay Mini-Program",
    )

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(
        intro,
        "This archive contains handwritten and structured evidence of offline teamwork across five key meetings.",
        size=12,
    )

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Project", "Dorm Takeaway Relay Mini-Program"),
        ("Group", "Group 15"),
        ("Archive Scope", "Five key meetings, planning evidence, and teamwork records"),
        ("Live Demo", "QR code to be inserted before final printing"),
    ]
    for row, values in zip(meta.rows, meta_rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.42)
        write_cell(row.cells[0], values[0], bold=True, size=10)
        write_cell(row.cells[1], values[1], size=10)

    doc.add_paragraph()
    team_heading = doc.add_paragraph()
    write_run(team_heading, "Team Roles", bold=True, size=13)

    team = doc.add_table(rows=1, cols=2)
    team.style = "Table Grid"
    team.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(team.rows[0].cells[0], "Member", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(team.rows[0].cells[1], "Role", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, role in TEAM:
        row = team.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        write_cell(row.cells[0], name, size=10)
        write_cell(row.cells[1], role, size=10)

    doc.add_paragraph()
    qr_heading = doc.add_paragraph()
    write_run(qr_heading, "QR Code Placement", bold=True, size=13)

    qr_box = doc.add_table(rows=1, cols=1)
    qr_box.style = "Table Grid"
    qr_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    qr_box.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    qr_box.rows[0].height = Inches(1.8)
    write_cell(
        qr_box.rows[0].cells[0],
        "Insert final deployed QR code here before printing.\nRecommended: center the code and keep one short URL backup below it.",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.add_paragraph()
    close = doc.add_paragraph()
    close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run(
        close,
        "Prepared for the teamwork defense as a complete archive pack, not as separate loose notes.",
        italic=True,
        size=11,
    )

    path = OUTPUT_DIR / "Archive_Cover_Page.docx"
    doc.save(path)
    return path


def build_index() -> Path:
    doc = Document()
    set_defaults(doc)
    add_center_title(
        doc,
        "Archive Index",
        "Teacher-facing navigation page for the meeting evidence pack",
    )

    intro = doc.add_paragraph()
    write_run(
        intro,
        "Recommended reading order: Meeting 1, Meeting 2, Meeting 5, Meeting 4, Meeting 3.",
        size=11,
    )

    sections = [
        ("1", "Problem and Scope", "Problem domain sheet and scoping planning sheet"),
        ("2", "Meeting 1", "Project selection and team formation"),
        ("3", "Meeting 2", "Tutor-feedback pivot and re-scoping"),
        ("4", "Meeting 3", "Integration and issue review"),
        ("5", "Meeting 4", "Lifecycle completion and verification"),
        ("6", "Meeting 5", "Final deployment, rehearsal, and defense preparation"),
        ("7", "Leadership and Summary", "Team leader report and team meeting summary"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Section", "Title", "What It Proves"]):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for no, title, purpose in sections:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.45)
        write_cell(row.cells[0], no, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(row.cells[1], title, size=10)
        write_cell(row.cells[2], purpose, size=10)

    doc.add_paragraph()
    prompt = doc.add_paragraph()
    write_run(prompt, "Fast question guide", bold=True, size=13)

    bullets = [
        "If asked about topic choice or early collaboration: open Meeting 1.",
        "If asked about why the team changed direction: open Meeting 2.",
        "If asked about honest progress checks: open Meeting 3 and Meeting 4.",
        "If asked about teamwork control before defense: open Meeting 5.",
        "If asked who did what: ask each member to show their own handwritten packet.",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        write_run(p, text, size=10.5)

    doc.add_section(WD_SECTION.NEW_PAGE)
    role_page = doc.add_paragraph()
    write_run(role_page, "Who Holds Which Folder", bold=True, size=14)

    roles = [
        ("Jiacheng Tao", "Main Archive Pack and Team Leader pages"),
        ("Pengyang Liu", "Personal packet and spare front matter copy"),
        ("Sicheng Wu", "Personal packet and Reference Pack"),
        ("Xiaofeng Dang", "Personal packet"),
        ("Ruilin Ma", "Personal packet"),
    ]

    role_table = doc.add_table(rows=1, cols=2)
    role_table.style = "Table Grid"
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(role_table.rows[0].cells[0], "Member", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(role_table.rows[0].cells[1], "Folder Responsibility", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, responsibility in roles:
        row = role_table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.4)
        write_cell(row.cells[0], name, size=10)
        write_cell(row.cells[1], responsibility, size=10)

    path = OUTPUT_DIR / "Archive_Index.docx"
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [build_cover(), build_index()]
    print("Created files:")
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
