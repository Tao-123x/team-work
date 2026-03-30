from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials" / "week1"

TEAM = [
    (
        "Jiacheng Tao",
        "Team Leader",
        "TJC-FlowRead",
        "Interest-driven English reading app with AI translation and a hover dictionary.",
    ),
    (
        "Pengyang Liu",
        "Product Owner",
        "LPY-Campus Idle Book Circulation Platform",
        "Campus second-hand textbook circulation platform with reuse and trust mechanisms.",
    ),
    (
        "Sicheng Wu",
        "Quality Controller",
        "WSC-Campus Efficiency WeChat Bot",
        "Campus efficiency assistant bot for reminders, notices, and student task support.",
    ),
    (
        "Xiaofeng Dang",
        "Team Member",
        "Weather Care Push System",
        "Automated WeChat weather care push with reminders, suggestions, and emotional content.",
    ),
    (
        "Ruilin Ma",
        "Team Member",
        "MRL-Word Memorization Mini-Program",
        "Vocabulary memorization mini-program focused on daily review and word learning.",
    ),
]

POST_CRITERIA_NOTES = [
    (
        "Practicality",
        "Does the idea solve a clear real-world problem?",
        "The team preferred ideas that addressed frequent student or daily-life pain points instead of abstract concepts.",
    ),
    (
        "Feasibility",
        "Can we deliver a stable MVP in time?",
        "The group leaned toward a direction that could be scoped quickly and implemented with a controlled technical path.",
    ),
    (
        "Presentation Potential",
        "Can we explain and demonstrate it clearly?",
        "Because the final defense is only 15 minutes, the project needed one simple story and a manageable demo flow.",
    ),
    (
        "Teamwork Suitability",
        "Can the work be divided meaningfully across roles?",
        "The team cared about whether product, quality, development, and leadership work could all be evidenced clearly.",
    ),
]

POST_DISCUSSION_SUMMARY = [
    (
        "Ideas that gained support",
        "Weather Care Push System and FlowRead received the strongest attention because both had a clear story and visible user value.",
    ),
    (
        "Concerns or disagreements",
        "Broader platform ideas looked interesting, but some members felt they would be harder to reduce into a stable classroom MVP.",
    ),
    (
        "Reason for the initial Week 1 choice",
        "The group initially preferred Weather Care Push System because it seemed practical, easy to explain, and realistic to start quickly.",
    ),
]

POST_MEMBER_NOTES = {
    "Jiacheng Tao": [
        "I paid most attention to whether the project could be managed clearly as a team instead of becoming too broad.",
        "I suggested that we should prefer an idea with a simple demo path and clear role ownership.",
        "I agreed that presentation value mattered from the first week, not only technical novelty.",
        "I left the meeting with responsibility for coordinating the next planning step and tool setup.",
    ],
    "Pengyang Liu": [
        "I focused on which proposal solved the clearest user pain point and could be reduced into a narrow MVP.",
        "I argued that a project with too many features would weaken our final product story.",
        "I agreed that the initial direction needed a practical user case and not only an interesting concept.",
        "I took away the task of clarifying scope and backlog priorities for the next stage.",
    ],
    "Sicheng Wu": [
        "I paid attention to which idea could be verified with clear evidence and realistic acceptance criteria.",
        "I noted that some options might create risk because they were harder to test or too dependent on unstable conditions.",
        "I agreed that the team should keep records from the very beginning if we wanted strong teamwork marks later.",
        "I left the meeting needing to think early about checklists, risk points, and review evidence.",
    ],
    "Xiaofeng Dang": [
        "I focused on which idea had the cleanest user-facing flow and the best demo potential.",
        "I thought visual clarity was important because the audience should understand the project quickly.",
        "I agreed that an idea with a stronger interaction story would be easier to present confidently.",
        "I took away the need to think about page flow and what a future prototype might need to show.",
    ],
    "Ruilin Ma": [
        "I focused on which project had the most controllable technical path and the clearest logic behind the user flow.",
        "I noted that some proposals looked attractive but might become harder to structure on the backend.",
        "I agreed that a simpler logic chain would help us build a more stable MVP under time pressure.",
        "I left the meeting thinking about how the selected direction could later translate into routes, state, and data handling.",
    ],
}

POST_TL_NOTES = [
    "The meeting compared five separate proposals before any initial direction was selected.",
    "The team did not choose based on personal preference alone; practicality, feasibility, presentation clarity, and teamwork value were all considered.",
    "Weather Care Push System was selected in Week 1 because it looked more presentation-friendly and easier to scope at that stage.",
    "The meeting also created the first shared structure for roles and tool use, which became the team's early collaboration baseline.",
]

DECISION_SUMMARY = [
    ("Initial Week 1 Selection", "Weather Care Push System"),
    ("Why It Was Selected", "The team saw it as practical, easier to scope, and presentation-friendly at that stage."),
    ("Early Team Setup", "Roles and initial tools were confirmed after the discussion."),
]

FOLLOW_UP_SUMMARY = [
    "Product Owner to draft the first backlog and clarify scope.",
    "Team Leader to arrange the next planning meeting and coordinate responsibilities.",
    "Quality Controller to prepare the first checklist and risk notes.",
    "Developers to begin technical route investigation for the selected direction.",
]

THEME = {
    "navy": RGBColor(31, 59, 115),
    "muted": RGBColor(88, 96, 112),
}


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    color=None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.color.rgb = THEME["navy"]
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = THEME["muted"]
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r.font.color.rgb = THEME["navy"]
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_meta(doc: Document, status_label: str, pre: bool) -> None:
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if pre:
        rows = [
            ("Week", "Week 1", "Document Status", status_label),
            ("Meeting Type", "Project Selection and Team Formation", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "__________________", "Location", "__________________"),
            ("Attendees", "All group members", "Recorder", "__________________"),
        ]
    else:
        rows = [
            ("Week", "Week 1", "Document Status", status_label),
            ("Meeting Type", "Project Selection and Team Formation", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "Use the actual Week 1 meeting time from the handwritten archive copy.", "Location", "Use the actual meeting location from the handwritten archive copy."),
            ("Attendees", "All group members", "Recorder", "Use the name written on the paper archive copy."),
        ]
    for row, values in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        for i, value in enumerate(values):
            write_cell(
                row.cells[i],
                value,
                bold=i in (0, 2),
                size=10,
                color=THEME["navy"] if i in (0, 2) else None,
            )
            if i in (0, 2):
                shade(row.cells[i], "EEF3FC")


def add_projects(doc: Document, with_notes: bool) -> None:
    cols = 5 if with_notes else 4
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Member", "Role", "Recommended Project", "Core Idea"]
    if with_notes:
        headers.append("Discussion Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=THEME["navy"])
        shade(cell, "DCE6F8")

    for member in TEAM:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.75 if with_notes else 0.45)
        values = list(member)
        if with_notes:
            values.append("\n" * 5)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_criteria(doc: Document, pre: bool) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Criteria", "Purpose", "Meeting Record"]
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=THEME["navy"])
        shade(cell, "DCE6F8")

    rows = [
        ("Practicality", "Does the idea solve a clear real-world problem?"),
        ("Feasibility", "Can we deliver a stable MVP in time?"),
        ("Presentation Potential", "Can we explain and demonstrate it clearly?"),
        ("Teamwork Suitability", "Can the work be divided meaningfully across roles?"),
    ]
    for idx, (label, purpose) in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.75 if pre else 0.58)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], purpose, size=10)
        if pre:
            write_cell(row.cells[2], "\n" * 5, size=10)
        else:
            write_cell(row.cells[2], POST_CRITERIA_NOTES[idx][2], size=10)


def add_large_notes_table(doc: Document, title: str, rows: list[str], height: float) -> None:
    add_heading(doc, title)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, label in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(height)
        write_cell(row.cells[0], label, bold=True, size=10, color=THEME["navy"])
        shade(row.cells[0], "F7F9FC")
        write_cell(row.cells[1], "\n" * 7, size=10)


def add_post_decisions(doc: Document) -> None:
    add_heading(doc, "Meeting Decisions")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (left, right) in zip(table.rows, DECISION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.45)
        write_cell(row.cells[0], left, bold=True, size=10, color=THEME["navy"])
        shade(row.cells[0], "EEF3FC")
        write_cell(row.cells[1], right, size=10)

    add_heading(doc, "Week 2 Follow-Up Actions")
    for item in FOLLOW_UP_SUMMARY:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Post-Meeting Discussion Summary")
    table = doc.add_table(rows=len(POST_DISCUSSION_SUMMARY), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DISCUSSION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.78)
        write_cell(row.cells[0], label, bold=True, size=10, color=THEME["navy"])
        shade(row.cells[0], "F7F9FC")
        write_cell(row.cells[1], note, size=10)


def add_pre_body(doc: Document) -> None:
    add_title(
        doc,
        "Week 1 Project Selection and Team Formation Meeting",
        "Pre-Meeting Discussion Base",
    )
    add_meta(doc, "Pre-meeting", pre=True)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Compare the five candidate project ideas brought by team members.",
        "Discuss which option should become the initial team direction.",
        "Start building shared role and workflow awareness.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Candidate Project Ideas")
    add_projects(doc, with_notes=True)
    add_heading(doc, "Discussion Criteria")
    add_criteria(doc, pre=True)
    add_large_notes_table(
        doc,
        "Open Discussion Space",
        [
            "Ideas that gained support",
            "Concerns or disagreements",
            "Possible decision direction",
        ],
        1.1,
    )
    add_large_notes_table(
        doc,
        "Decision and Next Actions (leave blank before the meeting)",
        [
            "Final direction decided in the meeting",
            "Roles or tools confirmed",
            "Actions before the next meeting",
        ],
        1.0,
    )


def add_post_body(doc: Document) -> None:
    add_title(
        doc,
        "Week 1 Project Selection and Team Formation Meeting",
        "Post-Meeting Archive Reference",
    )
    add_meta(doc, "Post-meeting", pre=False)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Compare the five candidate project ideas brought by team members.",
        "Choose one initial project direction for Week 1 planning.",
        "Confirm early role structure and collaboration tools.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Candidate Project Ideas")
    add_projects(doc, with_notes=False)
    add_heading(doc, "Discussion Criteria Summary")
    add_criteria(doc, pre=False)
    add_post_decisions(doc)


def add_individual_page(doc: Document, member: tuple[str, str, str, str], pre: bool) -> None:
    name, role, project, summary = member
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        f"Week 1 Individual Notes - {name}",
        "Pre-Meeting Personal Discussion Sheet" if pre else "Post-Meeting Personal Record",
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name", name, "Role", role),
        ("Recommended Project", project, "Core Idea", summary),
    ]
    for row, values in zip(table.rows, rows):
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10, color=THEME["navy"] if i in (0, 2) else None)
            if i in (0, 2):
                shade(row.cells[i], "EEF3FC")

    if pre:
        add_large_notes_table(
            doc,
            "Personal Handwritten Space",
            [
                "What I want to ask in the meeting",
                "What I think is strongest about my proposal",
                "What risks I already notice",
                "What I want to learn from other proposals",
            ],
            1.05,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Post-Meeting Personal Record")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "What I focused on during the meeting",
            "What I said or suggested",
            "What I agreed or disagreed with",
            "What task or next step I took away",
        ]
        for row, label, note in zip(table.rows, labels, POST_MEMBER_NOTES[name]):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.92)
            write_cell(row.cells[0], label, bold=True, size=10, color=THEME["navy"])
            shade(row.cells[0], "F7F9FC")
            write_cell(row.cells[1], note, size=10)


def add_tl_page(doc: Document, pre: bool) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        "Week 1 Team Leader Pre-Meeting Control Page" if pre else "Week 1 Team Leader Official Record",
        "Meeting setup and observation sheet" if pre else "Formal team archive version",
    )
    if pre:
        add_large_notes_table(
            doc,
            "Team Leader Writing Space",
            [
                "Meeting objectives to protect",
                "How to compare the five options fairly",
                "What decisions must be made before the meeting ends",
                "Potential role and workflow points to confirm",
            ],
            1.1,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Team Leader Official Summary")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "How the five proposals were compared",
            "Why the initial direction was selected",
            "How roles and tools were confirmed",
            "Leader reflection on the meeting quality",
        ]
        for row, label, note in zip(table.rows, labels, POST_TL_NOTES):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10, color=THEME["navy"])
            shade(row.cells[0], "F7F9FC")
            write_cell(row.cells[1], note, size=10)


def create_doc(filename: str, pre: bool, member: tuple[str, str, str, str] | None = None, tl: bool = False) -> Path:
    doc = Document()
    set_defaults(doc)
    if pre:
        add_pre_body(doc)
    else:
        add_post_body(doc)

    if member:
        add_individual_page(doc, member, pre=pre)
    if tl:
        add_tl_page(doc, pre=pre)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_doc("Week1_PreMeeting_Shared_Base.docx", pre=True),
        create_doc("Week1_PostMeeting_Shared_Record.docx", pre=False),
        create_doc("Week1_PreMeeting_Team_Leader_Page.docx", pre=True, tl=True),
        create_doc("Week1_PostMeeting_Team_Leader_Official_Record.docx", pre=False, tl=True),
    ]

    for member in TEAM:
        safe = member[0].replace(" ", "_")
        created.append(create_doc(f"Week1_PreMeeting_{safe}_Notes.docx", pre=True, member=member))
        created.append(create_doc(f"Week1_PostMeeting_{safe}_Notes.docx", pre=False, member=member))

    print("Created files:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
