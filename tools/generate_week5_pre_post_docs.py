from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials" / "week5"

TEAM = [
    ("Jiacheng Tao", "Team Leader"),
    ("Pengyang Liu", "Product Owner"),
    ("Sicheng Wu", "Quality Controller"),
    ("Xiaofeng Dang", "Team Member (Development)"),
    ("Ruilin Ma", "Team Member (Development)"),
]

WORKSTREAMS = [
    (
        "Deployment and access readiness",
        "Review whether the project can be opened reliably through the planned runtime setup and QR-code path.",
    ),
    (
        "Locked demo path",
        "Confirm one short, safe, and repeatable live flow for the final defense.",
    ),
    (
        "Rehearsal and speaking coordination",
        "Check whether the five speakers, timing, and handoff points now support a controlled 15-minute defense.",
    ),
    (
        "Evidence packaging",
        "Review handwritten materials, printed archive order, and backup plans for the presentation.",
    ),
]

POST_CRITERIA_NOTES = [
    (
        "Deployment Readiness",
        "Can the team open the product reliably in the live defense setting?",
        "The group treated access, QR-code entry, and runtime setup as part of the final meeting rather than something to improvise later.",
    ),
    (
        "Demo Control",
        "Is the live demo now short, stable, and easy to narrate?",
        "The team wanted one locked happy-path flow rather than a broad or risky walkthrough.",
    ),
    (
        "Evidence Readiness",
        "Are the printed materials and support documents ready to strengthen teamwork marks?",
        "The group reviewed how handwritten meeting papers and supporting records would appear during the defense.",
    ),
    (
        "Defense Coordination",
        "Do speaking order, timing, and contingency plans now work together?",
        "The meeting linked technical preparation with speaker coordination so the final presentation would feel managed rather than improvised.",
    ),
]

POST_DECISIONS = [
    ("Main Meeting Outcome", "The fifth and final key meeting locked the defense setup around one controlled demo path, one evidence pack, and one speaker sequence."),
    ("Deployment Outcome", "The team agreed that runtime access and QR-based usage had to be prepared as a presentation requirement, not a last-minute technical task."),
    ("Defense Outcome", "The group treated rehearsal, printed archive order, and backup plans as part of final readiness rather than optional extras."),
]

POST_FOLLOW_UP = [
    "Product Owner to keep the demo story narrow and aligned with the locked live flow.",
    "Team Leader to finalize speaker timing, meeting-paper order, and defense run sequence.",
    "Quality Controller to prepare the final pre-defense checks and backup response plan.",
    "Developers to avoid destabilizing the locked demo path and support the final deployment setup.",
]

POST_DISCUSSION_SUMMARY = [
    (
        "What this meeting finalized",
        "The project moved from being presentation-ready in principle to being organized for an actual live defense.",
    ),
    (
        "Why that mattered",
        "The team understood that strong marks would depend not only on the product, but on how confidently and coherently the evidence and demo were delivered.",
    ),
    (
        "How the project changed here",
        "This final meeting brought deployment, printed materials, rehearsal, and role coordination into one controlled presentation package.",
    ),
]

POST_MEMBER_NOTES = {
    "Jiacheng Tao": [
        "I focused on whether the team now had one fully controlled defense setup instead of several partial ideas about how to present.",
        "I pushed the discussion toward final risks such as access, timing, speaking handoff, and whether the evidence pack would support our strongest marks.",
        "I agreed that the final meeting had to lock decisions rather than reopen scope, because the defense was now close.",
        "I left the meeting responsible for coordinating the final run order, printed archive structure, and overall defense flow.",
    ],
    "Pengyang Liu": [
        "I focused on whether the final demo still told one simple and convincing product story from beginning to end.",
        "I argued that the audience should see only the most understandable route through the system, not every possible feature or exception.",
        "I agreed that the product story needed to stay narrower in the defense than in the full implementation history.",
        "I took away the task of protecting the clearest live flow and making sure the explanation stayed user-centered.",
    ],
    "Sicheng Wu": [
        "I focused on whether the final defense path was still supported by genuine checks, controlled constraints, and realistic backup plans.",
        "I highlighted that the last meeting should not only celebrate readiness but also identify what must not change before the defense.",
        "I agreed that good presentation depends on trustworthy evidence as much as on confident speaking.",
        "I left the meeting planning the final readiness checklist and the safest backup responses if something goes wrong live.",
    ],
    "Xiaofeng Dang": [
        "I focused on whether the final visible screens now supported a confident demo without awkward pauses or confusing jumps.",
        "I suggested that the live sequence should move only through the cleanest and most legible interactions.",
        "I agreed that visual clarity and speaking clarity had to reinforce each other in the final defense.",
        "I took away the need to protect the few visible screens that would matter most during the presentation.",
    ],
    "Ruilin Ma": [
        "I focused on whether the locked demo path still rested on a stable runtime state, role permissions, and repeatable preparation steps.",
        "I noted that final presentation shortcuts are only safe if they match the actual system behavior underneath.",
        "I agreed that the last meeting should protect stability more than ambition.",
        "I left the meeting clearer about keeping the runtime path safe and avoiding risky last-minute edits before the defense.",
    ],
}

POST_TL_NOTES = [
    "The fifth meeting worked as the final coordination point before defense preparation was fully locked.",
    "Its value was not only technical; it also aligned deployment, printed evidence, speaker timing, and demo control into one presentation plan.",
    "The meeting helped the team decide what to keep fixed and what not to touch before the defense.",
    "It therefore closed the project history as a five-meeting chain rather than opening a separate sixth meeting.",
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def write_cell(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(10.5)
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_meta(doc: Document, status_label: str, pre: bool) -> None:
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if pre:
        rows = [
            ("Week", "Week 5", "Document Status", status_label),
            ("Meeting Type", "Final Deployment, Rehearsal, and Defense Preparation", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "__________________", "Location", "__________________"),
            ("Attendees", "All group members", "Recorder", "__________________"),
        ]
    else:
        rows = [
            ("Week", "Week 5", "Document Status", status_label),
            ("Meeting Type", "Final Deployment, Rehearsal, and Defense Preparation", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "Use the actual Week 5 meeting time from the handwritten archive copy.", "Location", "Use the actual meeting location from the handwritten archive copy."),
            ("Attendees", "All group members", "Recorder", "Use the recorder written on the paper archive copy."),
        ]
    for row, values in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)


def add_participants(doc: Document, pre: bool) -> None:
    add_heading(doc, "Attendees and Roles")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Member", "Role"]
    if pre:
        headers.append("Personal Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for name, role in TEAM:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.52 if pre else 0.36)
        values = [name, role]
        if pre:
            values.append("\n" * 4)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_workstreams(doc: Document, pre: bool) -> None:
    add_heading(doc, "Workstreams Under Review")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Workstream", "Review Focus"]
    if pre:
        headers.append("Meeting Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for workstream, note in WORKSTREAMS:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.68 if pre else 0.45)
        values = [workstream, note]
        if pre:
            values.append("\n" * 5)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_criteria(doc: Document, pre: bool) -> None:
    add_heading(doc, "Review Criteria")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Criteria", "Purpose", "Meeting Record"]):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    criteria = [
        ("Demo Repeatability", "Can the team run the same stable presentation path again and again?"),
        ("Presentation Clarity", "Does the interface now make the core story easier for the audience to follow?"),
        ("Build Confidence", "Do current builds support the intended presentation platforms?"),
        ("Credible Constraints", "Does the system now handle key invalid actions in a believable way?"),
    ]
    for idx, (label, purpose) in enumerate(criteria):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.78 if pre else 0.56)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], purpose, size=10)
        if pre:
            write_cell(row.cells[2], "\n" * 5, size=10)
        else:
            write_cell(row.cells[2], POST_CRITERIA_NOTES[idx][2], size=10)


def add_large_notes_table(doc: Document, title: str, rows: list[str], height: float) -> None:
    add_heading(doc, title)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, label in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(height)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], "\n" * 7, size=10)


def add_post_summary(doc: Document) -> None:
    add_heading(doc, "Meeting Decisions")
    table = doc.add_table(rows=len(POST_DECISIONS), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DECISIONS):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.48)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)

    add_heading(doc, "Final Actions Before Defense")
    for item in POST_FOLLOW_UP:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Post-Meeting Discussion Summary")
    table = doc.add_table(rows=len(POST_DISCUSSION_SUMMARY), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DISCUSSION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.82)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)


def add_pre_body(doc: Document) -> None:
    add_title(doc, "Meeting 5 Final Deployment, Rehearsal, and Defense Preparation", "Pre-Meeting Discussion Base")
    add_meta(doc, "Pre-meeting", pre=True)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review whether the current build and deployment setup are ready for the live defense.",
        "Lock the final demo route, printed evidence order, and speaker handoff for the 15-minute presentation.",
        "Confirm backup plans so the team can still present confidently if a live step becomes unstable.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=True)
    add_workstreams(doc, pre=True)
    add_criteria(doc, pre=True)
    add_large_notes_table(
        doc,
        "Open Discussion Space",
        [
            "What now makes the final demo path safe and controllable",
            "What still feels risky or unclear in the live defense setup",
            "What must be protected or printed before presentation day",
        ],
        1.05,
    )
    add_large_notes_table(
        doc,
        "Decision and Next Actions (leave blank before the meeting)",
        [
            "Main final-defense conclusions agreed in the meeting",
            "Deployment, rehearsal, and archive priorities selected here",
            "Actions to complete before presentation day",
        ],
        1.0,
    )


def add_post_body(doc: Document) -> None:
    add_title(doc, "Meeting 5 Final Deployment, Rehearsal, and Defense Preparation", "Post-Meeting Archive Reference")
    add_meta(doc, "Post-meeting", pre=False)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review whether the current build and deployment setup are ready for the live defense.",
        "Lock the final demo route, printed evidence order, and speaker handoff for the 15-minute presentation.",
        "Confirm backup plans so the team can still present confidently if a live step becomes unstable.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=False)
    add_workstreams(doc, pre=False)
    add_criteria(doc, pre=False)
    add_post_summary(doc)


def add_individual_page(doc: Document, member: tuple[str, str], pre: bool) -> None:
    name, role = member
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        f"Week 5 Individual Notes - {name}",
        "Pre-Meeting Personal Discussion Sheet" if pre else "Post-Meeting Personal Record",
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name", name, "Role", role),
        ("Meeting Focus", "Final deployment, rehearsal, and defense preparation", "Meeting", "Meeting 5"),
    ]
    for row, values in zip(table.rows, rows):
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)

    if pre:
        add_large_notes_table(
            doc,
            "Personal Handwritten Space",
            [
                "What I most want to stabilize before the defense",
                "What still looks risky in the live demo path",
                "What improvement matters most for audience understanding",
                "What I need clarified before the meeting ends",
            ],
            1.05,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Post-Meeting Personal Record")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "What I focused on during the presentation-readiness review",
            "What I said or suggested",
            "What I agreed or disagreed with",
            "What task or next step I took away",
        ]
        for row, label, note in zip(table.rows, labels, POST_MEMBER_NOTES[name]):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def add_tl_page(doc: Document, pre: bool) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        "Meeting 5 Team Leader Pre-Meeting Control Page" if pre else "Meeting 5 Team Leader Official Record",
        "Meeting setup and observation sheet" if pre else "Formal team archive version",
    )
    if pre:
        add_large_notes_table(
            doc,
            "Team Leader Writing Space",
            [
                "What presentation risks must be checked honestly in this meeting",
                "How to keep the group focused on stability rather than new scope",
                "What final defense decisions must be fixed before the meeting ends",
                "How to prepare the project for rehearsal and packaging",
            ],
            1.08,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Team Leader Official Summary")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "How final readiness was evaluated",
            "Why deployment and rehearsal became the main meeting topic",
            "How the team prepared to lock the defense setup",
            "Leader reflection on the meeting quality",
        ]
        for row, label, note in zip(table.rows, labels, POST_TL_NOTES):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def create_doc(filename: str, pre: bool, member: tuple[str, str] | None = None, tl: bool = False) -> Path:
    doc = Document()
    set_defaults(doc)
    if pre:
        add_pre_body(doc)
    else:
        add_post_body(doc)

    if member:
        add_individual_page(doc, member, pre=pre)
    if tl:
        add_tl_page(doc, pre=pre)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_doc("Week5_PreMeeting_Shared_Base.docx", pre=True),
        create_doc("Week5_PostMeeting_Shared_Record.docx", pre=False),
        create_doc("Week5_PreMeeting_Team_Leader_Page.docx", pre=True, tl=True),
        create_doc("Week5_PostMeeting_Team_Leader_Official_Record.docx", pre=False, tl=True),
    ]

    for member in TEAM:
        safe = member[0].replace(" ", "_")
        created.append(create_doc(f"Week5_PreMeeting_{safe}_Notes.docx", pre=True, member=member))
        created.append(create_doc(f"Week5_PostMeeting_{safe}_Notes.docx", pre=False, member=member))

    print("Created files:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
