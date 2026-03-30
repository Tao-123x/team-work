from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials" / "week2"

TEAM = [
    ("Jiacheng Tao", "Team Leader"),
    ("Pengyang Liu", "Product Owner"),
    ("Sicheng Wu", "Quality Controller"),
    ("Xiaofeng Dang", "Team Member (Development)"),
    ("Ruilin Ma", "Team Member (Development)"),
]

PIVOT_OPTIONS = [
    (
        "Keep improving the Week 1 direction",
        "Would save some early effort, but did not respond well enough to tutor feedback.",
    ),
    (
        "Choose another broad campus platform idea",
        "Could look ambitious, but risked weak scope control and a confusing final demo.",
    ),
    (
        "Pivot to a dormitory takeaway relay mini-program",
        "Matched a practical student pain point and was easier to scope into a clear MVP.",
    ),
]

POST_CRITERIA_NOTES = [
    (
        "Response to Feedback",
        "Does the new direction answer the tutor's concerns?",
        "The team wanted a replacement topic that clearly solved the weaknesses identified after Week 1.",
    ),
    (
        "Student Relevance",
        "Does the project solve an obvious campus problem?",
        "A high-frequency dormitory pain point would be easier to justify in both the report and the defense.",
    ),
    (
        "MVP Clarity",
        "Can we define a small but complete workflow quickly?",
        "The new topic needed one clear user journey from posting a task to completing delivery.",
    ),
    (
        "Execution Readiness",
        "Can the team split the work and start immediately?",
        "The pivot only made sense if backlog reset, role alignment, and early development could happen in the same week.",
    ),
]

POST_DECISIONS = [
    ("Project Pivot", "The team officially pivoted to a dormitory takeaway relay mini-program."),
    ("Why This Topic", "It addressed a clear dormitory pain point and could be presented through a focused end-to-end workflow."),
    ("Immediate Setup", "The backlog, sprint direction, and role responsibilities were refreshed around the new topic."),
]

POST_FOLLOW_UP = [
    "Product Owner to clarify the new MVP scope, user stories, and backlog priorities.",
    "Team Leader to reset the working plan and coordinate the first implementation split.",
    "Quality Controller to rebuild the checklist, risk notes, and testing focus for the new topic.",
    "Developers to begin framework setup and early page or logic exploration for the relay workflow.",
]

POST_DISCUSSION_SUMMARY = [
    (
        "Main trigger for the pivot",
        "Tutor feedback made it clear that the original direction was not the best fit for the course goals and presentation expectations.",
    ),
    (
        "Why the new topic gained support",
        "The takeaway relay idea felt practical, student-facing, and easier to explain with one concrete scenario.",
    ),
    (
        "How the team reacted",
        "Instead of treating the change as a setback, the group used the meeting to reset scope, divide work, and keep momentum.",
    ),
]

POST_MEMBER_NOTES = {
    "Jiacheng Tao": [
        "I focused on how to stabilize the team quickly after the pivot instead of letting the change slow us down.",
        "I pushed for a direction that could support clear ownership, a simple demo path, and a manageable timeline.",
        "I agreed that the new topic had to improve both project fit and presentation control, not just novelty.",
        "I left the meeting with responsibility for coordinating the reset plan and making sure everyone had a clear next step.",
    ],
    "Pengyang Liu": [
        "I focused on redefining the user problem so the new topic would have a narrower and stronger MVP boundary.",
        "I argued that the relay workflow should stay simple: post, accept, pick up, deliver, complete.",
        "I agreed that the pivot would only help if we avoided feature expansion and kept the backlog disciplined.",
        "I took away the task of rewriting the backlog and clarifying the core user stories for the new product direction.",
    ],
    "Sicheng Wu": [
        "I focused on whether the new topic could be tested more clearly than the rejected proposal.",
        "I pointed out that the pivot needed a fresh checklist, new acceptance criteria, and updated risk notes.",
        "I agreed that changing the topic was acceptable only if the team could preserve evidence and quality control.",
        "I left the meeting planning to rebuild the testing focus around roles, state flow, and order handling.",
    ],
    "Xiaofeng Dang": [
        "I focused on whether the new relay idea could produce a more intuitive and understandable user interface.",
        "I suggested that a two-sided flow between requester and helper would create a clearer demo experience.",
        "I agreed that the topic was stronger because the audience could understand the value in a few seconds.",
        "I took away the need to think about page flow, task posting, and order-list interaction for the new design.",
    ],
    "Ruilin Ma": [
        "I focused on whether the new topic had a cleaner backend logic chain than the rejected proposal.",
        "I noted that order states and role transitions could give the system a clearer technical structure.",
        "I agreed that the relay workflow would be easier to implement in a stable way under limited time.",
        "I left the meeting thinking about routes, order data, and how to enforce the basic task lifecycle.",
    ],
}

POST_TL_NOTES = [
    "The Week 2 meeting was a controlled pivot rather than an unstructured topic change.",
    "The team used tutor feedback as a decision input and then compared replacement directions against scope, clarity, and readiness.",
    "The dormitory takeaway relay topic was selected because it gave the group a clearer problem statement and a stronger demo path.",
    "Before the meeting ended, responsibilities and next actions were reset so that the project could restart without losing momentum.",
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def write_cell(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(10.5)
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_meta(doc: Document, status_label: str, pre: bool) -> None:
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if pre:
        rows = [
            ("Week", "Week 2", "Document Status", status_label),
            ("Meeting Type", "Project Pivot and Re-Scoping", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "__________________", "Location", "__________________"),
            ("Attendees", "All group members", "Recorder", "__________________"),
        ]
    else:
        rows = [
            ("Week", "Week 2", "Document Status", status_label),
            ("Meeting Type", "Project Pivot and Re-Scoping", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "Use the actual Week 2 meeting time from the handwritten archive copy.", "Location", "Use the actual meeting location from the handwritten archive copy."),
            ("Attendees", "All group members", "Recorder", "Use the recorder written on the paper archive copy."),
        ]
    for row, values in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)


def add_participants(doc: Document, pre: bool) -> None:
    add_heading(doc, "Attendees and Roles")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Member", "Role"]
    if pre:
        headers.append("Personal Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for name, role in TEAM:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.52 if pre else 0.36)
        values = [name, role]
        if pre:
            values.append("\n" * 4)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_pivot_options(doc: Document, pre: bool) -> None:
    add_heading(doc, "Pivot Options Under Discussion")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Option", "Reasoning"]
    if pre:
        headers.append("Meeting Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for option, note in PIVOT_OPTIONS:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.68 if pre else 0.45)
        values = [option, note]
        if pre:
            values.append("\n" * 5)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_criteria(doc: Document, pre: bool) -> None:
    add_heading(doc, "Pivot Decision Criteria")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Criteria", "Purpose", "Meeting Record"]):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    criteria = [
        ("Response to Feedback", "Does the new direction answer the tutor's concerns?"),
        ("Student Relevance", "Does the project solve an obvious campus problem?"),
        ("MVP Clarity", "Can we define a small but complete workflow quickly?"),
        ("Execution Readiness", "Can the team split the work and start immediately?"),
    ]
    for idx, (label, purpose) in enumerate(criteria):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.78 if pre else 0.56)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], purpose, size=10)
        if pre:
            write_cell(row.cells[2], "\n" * 5, size=10)
        else:
            write_cell(row.cells[2], POST_CRITERIA_NOTES[idx][2], size=10)


def add_large_notes_table(doc: Document, title: str, rows: list[str], height: float) -> None:
    add_heading(doc, title)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, label in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(height)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], "\n" * 7, size=10)


def add_post_summary(doc: Document) -> None:
    add_heading(doc, "Meeting Decisions")
    table = doc.add_table(rows=len(POST_DECISIONS), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DECISIONS):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.48)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)

    add_heading(doc, "Week 3 Follow-Up Actions")
    for item in POST_FOLLOW_UP:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Post-Meeting Discussion Summary")
    table = doc.add_table(rows=len(POST_DISCUSSION_SUMMARY), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DISCUSSION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.82)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)


def add_pre_body(doc: Document) -> None:
    add_title(doc, "Week 2 Project Pivot and Re-Scoping Meeting", "Pre-Meeting Discussion Base")
    add_meta(doc, "Pre-meeting", pre=True)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review tutor feedback after Week 1 and decide whether the original direction should be changed.",
        "Compare replacement options and agree on a stronger project direction.",
        "Reset scope, responsibilities, and immediate actions for the next sprint.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=True)
    add_pivot_options(doc, pre=True)
    add_criteria(doc, pre=True)
    add_large_notes_table(
        doc,
        "Open Discussion Space",
        [
            "Tutor feedback points that changed our thinking",
            "Why the new topic gained support",
            "Concerns or disagreements during the pivot discussion",
        ],
        1.05,
    )
    add_large_notes_table(
        doc,
        "Decision and Next Actions (leave blank before the meeting)",
        [
            "Final project direction selected in the meeting",
            "New role or planning adjustments confirmed",
            "Actions to complete before Week 3 work begins",
        ],
        1.0,
    )


def add_post_body(doc: Document) -> None:
    add_title(doc, "Week 2 Project Pivot and Re-Scoping Meeting", "Post-Meeting Archive Reference")
    add_meta(doc, "Post-meeting", pre=False)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review tutor feedback after Week 1 and decide whether the original direction should be changed.",
        "Compare replacement options and agree on a stronger project direction.",
        "Reset scope, responsibilities, and immediate actions for the next sprint.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=False)
    add_pivot_options(doc, pre=False)
    add_criteria(doc, pre=False)
    add_post_summary(doc)


def add_individual_page(doc: Document, member: tuple[str, str], pre: bool) -> None:
    name, role = member
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        f"Week 2 Individual Notes - {name}",
        "Pre-Meeting Personal Discussion Sheet" if pre else "Post-Meeting Personal Record",
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name", name, "Role", role),
        ("Meeting Focus", "Project pivot and re-scoping", "Week", "Week 2"),
    ]
    for row, values in zip(table.rows, rows):
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)

    if pre:
        add_large_notes_table(
            doc,
            "Personal Handwritten Space",
            [
                "What I want to ask after hearing the tutor feedback",
                "What I think the new topic must improve",
                "What risk I notice in changing direction",
                "What I need to clarify before the meeting ends",
            ],
            1.05,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Post-Meeting Personal Record")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "What I focused on during the pivot meeting",
            "What I said or suggested",
            "What I agreed or disagreed with",
            "What task or next step I took away",
        ]
        for row, label, note in zip(table.rows, labels, POST_MEMBER_NOTES[name]):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def add_tl_page(doc: Document, pre: bool) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        "Week 2 Team Leader Pre-Meeting Control Page" if pre else "Week 2 Team Leader Official Record",
        "Meeting setup and observation sheet" if pre else "Formal team archive version",
    )
    if pre:
        add_large_notes_table(
            doc,
            "Team Leader Writing Space",
            [
                "What feedback points must be addressed in this meeting",
                "How to keep the pivot discussion focused and constructive",
                "What decisions must be confirmed before the meeting ends",
                "How to restart the team with clear responsibilities after the pivot",
            ],
            1.08,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Team Leader Official Summary")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "How the pivot decision was controlled",
            "Why the new direction was selected",
            "How responsibilities were reset after the change",
            "Leader reflection on the meeting quality",
        ]
        for row, label, note in zip(table.rows, labels, POST_TL_NOTES):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def create_doc(filename: str, pre: bool, member: tuple[str, str] | None = None, tl: bool = False) -> Path:
    doc = Document()
    set_defaults(doc)
    if pre:
        add_pre_body(doc)
    else:
        add_post_body(doc)

    if member:
        add_individual_page(doc, member, pre=pre)
    if tl:
        add_tl_page(doc, pre=pre)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_doc("Week2_PreMeeting_Shared_Base.docx", pre=True),
        create_doc("Week2_PostMeeting_Shared_Record.docx", pre=False),
        create_doc("Week2_PreMeeting_Team_Leader_Page.docx", pre=True, tl=True),
        create_doc("Week2_PostMeeting_Team_Leader_Official_Record.docx", pre=False, tl=True),
    ]

    for member in TEAM:
        safe = member[0].replace(" ", "_")
        created.append(create_doc(f"Week2_PreMeeting_{safe}_Notes.docx", pre=True, member=member))
        created.append(create_doc(f"Week2_PostMeeting_{safe}_Notes.docx", pre=False, member=member))

    print("Created files:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
