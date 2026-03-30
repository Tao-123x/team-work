from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "defense-materials" / "week3"

TEAM = [
    ("Jiacheng Tao", "Team Leader"),
    ("Pengyang Liu", "Product Owner"),
    ("Sicheng Wu", "Quality Controller"),
    ("Xiaofeng Dang", "Team Member (Development)"),
    ("Ruilin Ma", "Team Member (Development)"),
]

WORKSTREAMS = [
    (
        "User flow alignment",
        "Check whether requester and helper flows now match the scoped MVP agreed in Week 2.",
    ),
    (
        "Frontend and backend integration",
        "Review whether pages, APIs, and order states are starting to connect correctly.",
    ),
    (
        "Quality and bug review",
        "Identify early bugs, risky assumptions, and missing checks before the build grows further.",
    ),
    (
        "Week 4 priorities",
        "Decide what must be fixed, completed, or simplified in the next sprint.",
    ),
]

POST_CRITERIA_NOTES = [
    (
        "Workflow Completeness",
        "Does the relay process already form a believable end-to-end MVP?",
        "The team checked whether posting, accepting, and delivery state changes were now visible as one joined-up flow.",
    ),
    (
        "Integration Stability",
        "Are separate pieces beginning to work together reliably?",
        "The meeting focused on whether frontend pages, backend logic, and data assumptions still matched after initial implementation.",
    ),
    (
        "Risk Visibility",
        "Have key bugs and weak points already been surfaced?",
        "The team wanted problems to appear early enough to be fixed before the final defense stage.",
    ),
    (
        "Priority Control",
        "Can we choose the next sprint work without expanding scope?",
        "The group deliberately filtered issues so Week 4 would focus on critical integration and polish rather than feature drift.",
    ),
]

POST_DECISIONS = [
    ("Meeting Focus", "The group treated Week 3 as an integration and issue-review checkpoint rather than a new planning reset."),
    ("Main Outcome", "The team agreed that the relay workflow was taking shape, but some integration gaps and usability risks still needed attention."),
    ("Next Sprint Direction", "Week 4 would prioritize fixing flow inconsistencies, strengthening core interactions, and improving build stability."),
]

POST_FOLLOW_UP = [
    "Product Owner to refine the MVP boundary based on what was actually working after early integration.",
    "Team Leader to track unresolved issues and turn them into a focused Week 4 action list.",
    "Quality Controller to convert observed bugs and weak points into a clearer testing checklist.",
    "Developers to fix critical integration problems and strengthen the main user path before adding polish.",
]

POST_DISCUSSION_SUMMARY = [
    (
        "What looked promising",
        "The group could already see a recognizable requester-helper workflow, which helped confirm that the Week 2 pivot was workable.",
    ),
    (
        "What still felt weak",
        "Some implementation details, handoffs, or interaction points were not yet smooth enough for a confident demo path.",
    ),
    (
        "Why this meeting mattered",
        "It turned scattered progress into a shared picture of what was ready, what was risky, and what had to be fixed next.",
    ),
]

POST_MEMBER_NOTES = {
    "Jiacheng Tao": [
        "I focused on whether the team was still moving in one direction or whether implementation work had started to drift apart.",
        "I pushed the discussion toward concrete blockers instead of vague impressions so we could leave with a usable next-step list.",
        "I agreed that the project was progressing, but only if we kept Week 4 tightly focused on the main relay path.",
        "I left the meeting responsible for tracking the issue list and making sure responsibilities were clear before the next sprint.",
    ],
    "Pengyang Liu": [
        "I focused on whether the current build still matched the MVP story we had agreed after the pivot.",
        "I suggested that some weaker or distracting details should wait until the main workflow became smoother.",
        "I agreed that usability problems were now just as important as missing logic, because both affect the final demo.",
        "I took away the task of sharpening priorities so the next sprint would support one clear user journey.",
    ],
    "Sicheng Wu": [
        "I focused on whether the team had already exposed the most important risks before the system became harder to change.",
        "I highlighted bugs, weak checks, and places where the current workflow could confuse users or break expected behavior.",
        "I agreed that an honest issue review at this stage would protect both product quality and presentation confidence later.",
        "I left the meeting planning to formalize the observed issues into a clearer test and review checklist.",
    ],
    "Xiaofeng Dang": [
        "I focused on whether the interface already told a clean enough story for the audience to understand the relay process quickly.",
        "I suggested that some page flow and interaction details still needed simplification to reduce confusion.",
        "I agreed that design clarity and implementation stability had to improve together, not separately.",
        "I took away the need to adjust important screens and interaction steps around the main task flow.",
    ],
    "Ruilin Ma": [
        "I focused on whether the underlying logic and state transitions were becoming coherent as separate pieces were connected.",
        "I pointed out that weak integration now would become a much bigger problem if ignored for another sprint.",
        "I agreed that the team should strengthen the core data and state path before spending time on minor extras.",
        "I left the meeting with a clearer sense of which logic and integration points needed immediate attention.",
    ],
}

POST_TL_NOTES = [
    "Week 3 worked as a realistic project checkpoint: the team stopped to compare actual build progress against the intended MVP.",
    "The meeting was useful because it made integration gaps visible before the project moved too close to the final defense.",
    "Instead of broadening the scope, the group chose to tighten Week 4 around critical issues in the main relay flow.",
    "The meeting also helped keep every role aligned around the same short list of risks, fixes, and priorities.",
]


def set_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)


def write_cell(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.italic = True
    sr.font.name = "Times New Roman"
    sr.font.size = Pt(10.5)
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_meta(doc: Document, status_label: str, pre: bool) -> None:
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if pre:
        rows = [
            ("Week", "Week 3", "Document Status", status_label),
            ("Meeting Type", "Integration and Issue Review", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "__________________", "Location", "__________________"),
            ("Attendees", "All group members", "Recorder", "__________________"),
        ]
    else:
        rows = [
            ("Week", "Week 3", "Document Status", status_label),
            ("Meeting Type", "Integration and Issue Review", "Suggested Mode", "Physical or offline discussion"),
            ("Date and Time", "Use the actual Week 3 meeting time from the handwritten archive copy.", "Location", "Use the actual meeting location from the handwritten archive copy."),
            ("Attendees", "All group members", "Recorder", "Use the recorder written on the paper archive copy."),
        ]
    for row, values in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.36)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)


def add_participants(doc: Document, pre: bool) -> None:
    add_heading(doc, "Attendees and Roles")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Member", "Role"]
    if pre:
        headers.append("Personal Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for name, role in TEAM:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.52 if pre else 0.36)
        values = [name, role]
        if pre:
            values.append("\n" * 4)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_workstreams(doc: Document, pre: bool) -> None:
    add_heading(doc, "Workstreams Under Review")
    table = doc.add_table(rows=1, cols=3 if pre else 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Workstream", "Review Focus"]
    if pre:
        headers.append("Meeting Notes")
    for cell, text in zip(table.rows[0].cells, headers):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for workstream, note in WORKSTREAMS:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.68 if pre else 0.45)
        values = [workstream, note]
        if pre:
            values.append("\n" * 5)
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, size=10)


def add_criteria(doc: Document, pre: bool) -> None:
    add_heading(doc, "Review Criteria")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ["Criteria", "Purpose", "Meeting Record"]):
        write_cell(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    criteria = [
        ("Workflow Completeness", "Does the relay process already form a believable end-to-end MVP?"),
        ("Integration Stability", "Are separate pieces beginning to work together reliably?"),
        ("Risk Visibility", "Have key bugs and weak points already been surfaced?"),
        ("Priority Control", "Can we choose the next sprint work without expanding scope?"),
    ]
    for idx, (label, purpose) in enumerate(criteria):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.78 if pre else 0.56)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], purpose, size=10)
        if pre:
            write_cell(row.cells[2], "\n" * 5, size=10)
        else:
            write_cell(row.cells[2], POST_CRITERIA_NOTES[idx][2], size=10)


def add_large_notes_table(doc: Document, title: str, rows: list[str], height: float) -> None:
    add_heading(doc, title)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, label in zip(table.rows, rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(height)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], "\n" * 7, size=10)


def add_post_summary(doc: Document) -> None:
    add_heading(doc, "Meeting Decisions")
    table = doc.add_table(rows=len(POST_DECISIONS), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DECISIONS):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.48)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)

    add_heading(doc, "Week 4 Follow-Up Actions")
    for item in POST_FOLLOW_UP:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_heading(doc, "Post-Meeting Discussion Summary")
    table = doc.add_table(rows=len(POST_DISCUSSION_SUMMARY), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, note) in zip(table.rows, POST_DISCUSSION_SUMMARY):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.82)
        write_cell(row.cells[0], label, bold=True, size=10)
        write_cell(row.cells[1], note, size=10)


def add_pre_body(doc: Document) -> None:
    add_title(doc, "Week 3 Integration and Issue Review Meeting", "Pre-Meeting Discussion Base")
    add_meta(doc, "Pre-meeting", pre=True)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review what has actually been built since the Week 2 pivot.",
        "Identify integration gaps, usability concerns, and early bugs in the relay workflow.",
        "Decide which issues matter most before the next sprint begins.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=True)
    add_workstreams(doc, pre=True)
    add_criteria(doc, pre=True)
    add_large_notes_table(
        doc,
        "Open Discussion Space",
        [
            "What already works well enough to keep",
            "What feels broken, confusing, or risky",
            "What must be fixed before the next sprint ends",
        ],
        1.05,
    )
    add_large_notes_table(
        doc,
        "Decision and Next Actions (leave blank before the meeting)",
        [
            "Main integration issues agreed in the meeting",
            "Priority fixes selected for the next sprint",
            "Actions to complete before Week 4 review ends",
        ],
        1.0,
    )


def add_post_body(doc: Document) -> None:
    add_title(doc, "Week 3 Integration and Issue Review Meeting", "Post-Meeting Archive Reference")
    add_meta(doc, "Post-meeting", pre=False)

    add_heading(doc, "Meeting Purpose")
    for text in [
        "Review what has actually been built since the Week 2 pivot.",
        "Identify integration gaps, usability concerns, and early bugs in the relay workflow.",
        "Decide which issues matter most before the next sprint begins.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_participants(doc, pre=False)
    add_workstreams(doc, pre=False)
    add_criteria(doc, pre=False)
    add_post_summary(doc)


def add_individual_page(doc: Document, member: tuple[str, str], pre: bool) -> None:
    name, role = member
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        f"Week 3 Individual Notes - {name}",
        "Pre-Meeting Personal Discussion Sheet" if pre else "Post-Meeting Personal Record",
    )

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Name", name, "Role", role),
        ("Meeting Focus", "Integration and issue review", "Week", "Week 3"),
    ]
    for row, values in zip(table.rows, rows):
        for i, value in enumerate(values):
            write_cell(row.cells[i], value, bold=i in (0, 2), size=10)

    if pre:
        add_large_notes_table(
            doc,
            "Personal Handwritten Space",
            [
                "What I most want to check in the current build",
                "What issue or risk I already notice",
                "What I think is progressing well",
                "What I need clarified before the meeting ends",
            ],
            1.05,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Post-Meeting Personal Record")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "What I focused on during the integration review",
            "What I said or suggested",
            "What I agreed or disagreed with",
            "What task or next step I took away",
        ]
        for row, label, note in zip(table.rows, labels, POST_MEMBER_NOTES[name]):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def add_tl_page(doc: Document, pre: bool) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_title(
        doc,
        "Week 3 Team Leader Pre-Meeting Control Page" if pre else "Week 3 Team Leader Official Record",
        "Meeting setup and observation sheet" if pre else "Formal team archive version",
    )
    if pre:
        add_large_notes_table(
            doc,
            "Team Leader Writing Space",
            [
                "What build progress must be checked honestly in this meeting",
                "How to keep the review focused on real issues instead of excuses",
                "What priorities must be decided before the meeting ends",
                "How to prevent the team from expanding scope after finding problems",
            ],
            1.08,
        )
        sign = doc.add_table(rows=1, cols=3)
        sign.style = "Table Grid"
        for cell, text in zip(
            sign.rows[0].cells,
            [
                "Date: __________________",
                "Initials or Signature: __________________",
                "Location note: __________________",
            ],
        ):
            write_cell(cell, text, size=10)
    else:
        add_heading(doc, "Team Leader Official Summary")
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            "How Week 3 progress was evaluated",
            "Why certain issues became the next sprint priorities",
            "How the team stayed aligned after surfacing problems",
            "Leader reflection on the meeting quality",
        ]
        for row, label, note in zip(table.rows, labels, POST_TL_NOTES):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.94)
            write_cell(row.cells[0], label, bold=True, size=10)
            write_cell(row.cells[1], note, size=10)


def create_doc(filename: str, pre: bool, member: tuple[str, str] | None = None, tl: bool = False) -> Path:
    doc = Document()
    set_defaults(doc)
    if pre:
        add_pre_body(doc)
    else:
        add_post_body(doc)

    if member:
        add_individual_page(doc, member, pre=pre)
    if tl:
        add_tl_page(doc, pre=pre)

    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    created = [
        create_doc("Week3_PreMeeting_Shared_Base.docx", pre=True),
        create_doc("Week3_PostMeeting_Shared_Record.docx", pre=False),
        create_doc("Week3_PreMeeting_Team_Leader_Page.docx", pre=True, tl=True),
        create_doc("Week3_PostMeeting_Team_Leader_Official_Record.docx", pre=False, tl=True),
    ]

    for member in TEAM:
        safe = member[0].replace(" ", "_")
        created.append(create_doc(f"Week3_PreMeeting_{safe}_Notes.docx", pre=True, member=member))
        created.append(create_doc(f"Week3_PostMeeting_{safe}_Notes.docx", pre=False, member=member))

    print("Created files:")
    for path in created:
        print(path)


if __name__ == "__main__":
    main()
