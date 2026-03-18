from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
DOWNLOADS = Path.home() / "Downloads"

TEAM = [
    ("Team Leader", "Jiacheng Tao", "21907842"),
    ("Product Owner", "Pengyang Liu", "21907860"),
    ("Quality Controller", "Sicheng Wu", "21907856"),
    ("Team Member", "Xiaofeng Dang", "21907747"),
    ("Team Member", "Ruilin Ma", "21907748"),
    ("Team Member", "", ""),
]

GROUP_NAME = "Group 15"
PROJECT_NAME = "Dormitory Food Delivery Relay Mini-Program"


def set_paragraph(paragraph, text: str) -> None:
    paragraph.text = text


def set_cell(cell, text: str) -> None:
    cell.text = text


def set_row_height(row, inches: float) -> None:
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Inches(inches)


def insert_page_break_before(table) -> None:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    table._element.addprevious(paragraph)


def fill_team_table(table) -> None:
    set_cell(table.rows[2].cells[0], "GROUP NAME:\n" + GROUP_NAME)
    for index, (role, name, student_id) in enumerate(TEAM, start=6):
        row = table.rows[index]
        set_cell(row.cells[0], role)
        set_cell(row.cells[1], name)
        set_cell(row.cells[3], student_id)


def create_problem_domain_doc() -> Path:
    source = DOWNLOADS / "Project Problem Domain Activity Sheet.docx"
    target = ROOT / "Group 15_Project Problem Domain Activity Sheet.docx"
    doc = Document(str(source))

    fill_team_table(doc.tables[0])

    problem_table = doc.tables[1]
    set_cell(
        problem_table.rows[2].cells[1],
        "Students living in high-rise dormitories often cannot collect takeaway food from the ground-floor pickup area at the right time. This creates a last-meter delivery problem between the dorm entrance and the student’s room."
    )
    set_cell(
        problem_table.rows[3].cells[1],
        "We are doing this because the problem happens frequently in daily campus life. Students may be in class, ill, busy with study, or simply unable to go downstairs quickly. If the problem is not addressed, food gets cold, pickup areas become crowded, and the student experience is poor."
    )
    set_cell(
        problem_table.rows[4].cells[1],
        "We will address this by researching the dormitory use case, defining a clear MVP workflow, designing a mini-program interface for posting and accepting delivery requests, implementing the core order lifecycle, and evaluating the solution through walkthroughs, scenario testing, and team review."
    )

    doc.save(str(target))
    return target


def create_scoping_doc() -> Path:
    source = DOWNLOADS / "Project Scoping  Planning Activity Sheet.docx"
    target = ROOT / "Group 15_Project Scoping Planning Activity Sheet.docx"
    doc = Document(str(source))

    top = doc.tables[0]
    def pad(text: str) -> str:
        return " " + text

    set_cell(top.rows[2].cells[1], GROUP_NAME)
    set_cell(
        top.rows[5].cells[1],
        pad("Dorm takeaway relay.")
    )
    set_cell(
        top.rows[6].cells[1],
        pad("Students cannot always go downstairs in time.")
    )
    set_cell(
        top.rows[7].cells[1],
        pad("A real dormitory pain point.")
    )
    set_cell(
        top.rows[8].cells[1],
        pad("Build a simple peer-to-peer MVP.")
    )
    set_row_height(top.rows[5], 0.75)
    set_row_height(top.rows[6], 0.9)
    set_row_height(top.rows[7], 0.75)
    set_row_height(top.rows[8], 0.9)

    initial = doc.tables[1]
    set_cell(
        initial.rows[2].cells[1],
        pad("Students post and accept relay tasks.")
    )
    set_cell(initial.rows[3].cells[2], pad("Post request"))
    set_cell(initial.rows[4].cells[2], pad("Accept order"))
    set_cell(initial.rows[5].cells[2], pad("Confirm delivery"))
    set_row_height(initial.rows[2], 0.95)
    set_row_height(initial.rows[3], 0.75)
    set_row_height(initial.rows[4], 0.75)
    set_row_height(initial.rows[5], 0.85)

    solution = doc.tables[2]
    set_cell(
        solution.rows[2].cells[1],
        pad("Sprint 1: Confirm scope and first prototype.")
    )
    set_cell(
        solution.rows[3].cells[1],
        pad("Sprint 2: Build the main order flow.")
    )
    set_cell(
        solution.rows[4].cells[1],
        pad("Sprint 3: Add testing and presentation polish.")
    )
    set_row_height(solution.rows[2], 0.6)
    set_row_height(solution.rows[3], 0.6)
    set_row_height(solution.rows[4], 0.6)

    plan = doc.tables[3]
    set_cell(
        plan.rows[6].cells[1],
        pad("Finish the sprint with a validated direction and a working technical foundation.")
    )

    goals = [
        ("Goal 1:", "Confirm the problem definition, users, MVP scope, and acceptance criteria.", "8"),
        ("Goal 2:", "Set up the Trello backlog, sprint plan, and task ownership.", "8"),
        ("Goal 3:", "Design the page flow for hall, create-order, detail, and my-orders pages.", "12"),
        ("Goal 4:", "Set up the backend structure, data model, and API contracts.", "12"),
        ("Goal 5:", "Build the first prototype and prepare the QA checklist and risk register.", "10"),
    ]

    row_pairs = [(7, 8), (9, 10), (11, 12), (13, 14), (15, 16)]
    for (label_row, value_row), (_, objective, hours) in zip(row_pairs, goals):
        set_cell(plan.rows[label_row].cells[2], pad(objective))
        set_cell(plan.rows[label_row].cells[4], hours)
        set_cell(plan.rows[value_row].cells[2], "")
        set_cell(plan.rows[value_row].cells[4], "")

    set_cell(plan.rows[17].cells[4], "50")

    insert_page_break_before(doc.tables[1])
    insert_page_break_before(doc.tables[2])
    insert_page_break_before(doc.tables[3])

    doc.save(str(target))
    return target


def create_meeting_summary_doc() -> Path:
    source = DOWNLOADS / "Team Meeting Summary Template.docx"
    target = ROOT / "Group 15_Team Meeting Summary.docx"
    doc = Document(str(source))

    fill_team_table(doc.tables[0])

    meeting = doc.tables[1]
    set_cell(meeting.rows[3].cells[0], "Week 2 Meeting Record")
    set_cell(meeting.rows[4].cells[0], "Week 2")
    set_cell(meeting.rows[4].cells[1], "Not recorded")
    set_cell(meeting.rows[4].cells[2], "Not recorded")
    set_cell(
        meeting.rows[4].cells[3],
        "Reviewed instructor feedback on the rejected Week 1 proposal.\nAgreed to pivot from the weather-push concept to a dormitory food delivery relay mini-program.\nConfirmed core peer-to-peer order flow, updated the Trello backlog, and assigned Week 3 development and testing tasks."
    )
    set_cell(meeting.rows[4].cells[4], "All members")

    doc.save(str(target))
    return target


def create_team_leader_report_doc() -> Path:
    source = DOWNLOADS / "Team Leader Report Template.docx"
    target = ROOT / "Group 15_Team Leader Report.docx"
    doc = Document(str(source))

    fill_team_table(doc.tables[0])
    set_cell(doc.tables[0].rows[2].cells[0], "GROUP NAME:\n" + GROUP_NAME)

    progress = doc.tables[1]
    entries = {
        3: (
            "Initial proposal reviewed and challenged for low innovation. Team agreed that the first direction was not strong enough for the module.",
            "Identify a stronger real-world problem and redefine the project scope.",
            "Behind",
        ),
        4: (
            "The team selected a new project direction based on a real dormitory need and reset the backlog around the food delivery relay concept.",
            "Confirm the MVP scope, assign tasks, and prepare technical setup.",
            "Ontime",
        ),
        5: (
            "Core project planning was completed, including Trello backlog updates, role alignment, risk thinking, and initial framework research.",
            "Begin implementation of page flow, backend structure, and first integrated prototype.",
            "Ontime",
        ),
        6: (
            "The repo now includes the initial frontend pages, API structure, data model, and order lifecycle logic for the MVP.",
            "Complete testing evidence, polish edge cases, and continue integration work.",
            "Ontime",
        ),
    }
    for row_index, (done, pending, schedule) in entries.items():
        set_cell(progress.rows[row_index].cells[1], done)
        set_cell(progress.rows[row_index].cells[2], pending)
        set_cell(progress.rows[row_index].cells[3], schedule)

    engagement = doc.tables[2]
    set_cell(
        engagement.rows[3].cells[1],
        "The team reacted well to the project pivot. Members contributed to redefining the problem, and responsibilities remained clear despite the change in direction."
    )
    set_cell(
        engagement.rows[4].cells[1],
        "Engagement remained stable during implementation. Product, quality, and development activities were all represented, and team communication supported quick alignment."
    )

    issues = doc.tables[3]
    set_cell(
        issues.rows[3].cells[1],
        "The main issue was that the original proposal lacked sufficient innovation. This was resolved by re-scoping the project around a genuine dormitory pain point and updating the backlog immediately."
    )
    set_cell(
        issues.rows[4].cells[1],
        "A second issue was keeping planning, technical setup, and evaluation aligned after the pivot. This was handled through clearer sprint planning and a stronger focus on the MVP."
    )

    doc.save(str(target))
    return target


def main() -> None:
    outputs = [
        create_problem_domain_doc(),
        create_scoping_doc(),
        create_meeting_summary_doc(),
        create_team_leader_report_doc(),
    ]
    for path in outputs:
        print(f"Created: {path}")


if __name__ == "__main__":
    main()
